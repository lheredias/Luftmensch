from PyQt5.QtCore import (pyqtSignal,QThreadPool,pyqtSlot,QRunnable,QObject,Qt)
from PyQt5.QtWidgets import (QApplication, QMainWindow,QLabel,QFileDialog,QAction,
                             QProgressBar, QPushButton,QMessageBox,QLineEdit,QMenu,QComboBox,
                             QHBoxLayout,QStackedLayout,QGraphicsOpacityEffect,QAbstractButton,
                             QCheckBox,QVBoxLayout,QWidget,QListView)

from PyQt5.QtGui import (QIcon,QFont,QPixmap,QCursor,QPainter)
import time
import sys
import os
import getpass
import shutil
import fitz
from win32com import client
from docx import Document  
from webbrowser import open as op

# os.chdir(r'/Users/lenin/Documents/Python Scripts/Luftmensch')

if hasattr(Qt, 'AA_EnableHighDpiScaling'):
    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)

if hasattr(Qt, 'AA_UseHighDpiPixmaps'):
    QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    
    return os.path.join(base_path, relative_path)

icon=resource_path('finalicon.ico')
pic=resource_path('check small.png')
logo=resource_path('app name.png')
# Install Ghostscript and look for "bin" folder inside instalation directory.
# Copy all four files to your current working directory et voilà.
# Take these staps only if you're going to compile the file into an .exe.
# If not, just comment the "gs" definition line below.
gs=resource_path('gswin64c.exe')
# ======================================================================
donate_pic=resource_path('donate.png')
username=getpass.getuser()

years=[str(i) for i in range(2010,2021)]
years.reverse()
months=[str(i) for i in range(1,13)]
for i in range(len(months)):
    if len(months[i])==1:
        months[i]='0'+months[i]
months.reverse()
choices=['1. PDF to PDF/A',
         '2. PDF with Vertical A4 dimensions',
                 '3. Merge PDFs',
                 '4. Images to PDF',
                 '5. Compress PDF'] 

fontOne = QFont("Helvetica", 9)
fontTwo=QFont("Helvetica", 9)
fontThree=QFont('Consolas', 11)  #Done message font
fontFive=QFont('Consolas', 11) #Version font
# <codecell>
class WorkerSignalsOne(QObject):
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    
class JobRunnerOne(QRunnable):    
    signals = WorkerSignalsOne()
    
    def __init__(self,SaveAs,state):
        super().__init__()

        self.is_killed = False 
        self.SaveAs=SaveAs
        self.state=state
        
    @pyqtSlot()
    def is_opened(self):
        temp_filename=self.SaveAs[:-4]+' temp.pdf'
        if os.path.exists(self.SaveAs) == True:
            try:              
                os.rename(self.SaveAs,temp_filename)
                os.rename(temp_filename,self.SaveAs)               
                return False
            except PermissionError:
                return True
        else:
            return False
        
    def run(self):
       
        try:
            if self.is_opened() == True:
                self.signals.alert.emit('Error2')
            else:
                           
                done=False
                count=0
                while not done:
                    try:
                        backup=shutil.copy(self.SaveAs,self.SaveAs[:-4]+' sehr witzig.pdf')
                        done=True
                    except PermissionError:
                        print('Permission denied')
                        count+=1 
                        time.sleep(count)          
                time.sleep(1)
                DestDir=os.path.abspath(os.path.join(self.SaveAs, '..'))                
                tempDocx='\\tempword.docx'
                documento = Document()
                documento.save(DestDir+tempDocx)  
                try:
                    word = client.DispatchEx('Word.Application')
                except Exception:
                    self.signals.alert.emit('NoWord')
                else:
                    worddoc = word.Documents.Open(DestDir+tempDocx,ReadOnly = 1)                   
                    worddoc.SaveAs(self.SaveAs,FileFormat = 17)               
                    worddoc.Close(True)
                    word.Quit()
                    os.remove(DestDir+tempDocx)    
                    pdf=fitz.open(self.SaveAs)
                    opened_file=fitz.open(backup)
                    pdf.insertPDF(opened_file)
                    opened_file.close()
                    pdf.deletePage(0)
                    pdf.saveIncr()
                    pdf.close()               
                    time.sleep(1)      
                    done=False
                    count=0
                    while not done:
                        try:
                            os.remove(backup)
                            done=True
                        except PermissionError:
                            print('Permission denied')
                            count+=1 
                            time.sleep(count)
                    if self.state==2:
                        from subprocess import Popen 
                        Popen([self.SaveAs],shell=True)
                    time.sleep(1)
                    self.signals.finished.emit('Done')
        except Exception as e:      
               self.signals.alert.emit(str(type(e)))
    def kill(self):
        self.is_killed = True
          
class ActionsOne(QWidget):

    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.initUI()
        self.msg1='Choose a PDF file.'
        
    def initUI(self):
        
        self.style = QApplication.style()
        
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}")             
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
          
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))

        
        self.buttonTwo = QPushButton('Load PDF', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        # self.myTextBoxOne.setMaximumWidth(600)
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h1.addWidget(self.myTextBoxOne,4)
        
        # self.lineOne = QLabel('/'*250, self) 
        # self.lineOne.setMaximumWidth(800)
        # self.v1.addWidget(self.lineOne)
        
        self.CheckOne = QCheckBox('Open immediately', self)  
        self.CheckOne.setFont(fontTwo)
        self.CheckOne.setMinimumHeight(35)
        # self.CheckOne.setMaximumWidth(800)
        self.CheckOne.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}") 
        self.CheckOne.setChecked(True)
        self.v1.addWidget(self.CheckOne)
             
        # self.lineTwo = QLabel('/'*250, self)
        # self.lineTwo.setMaximumWidth(800)
        # self.v1.addWidget(self.lineTwo)     
        self.h2.addStretch()
        self.start = QPushButton('Run', self)
        self.start.setStyleSheet(self.style1)
        # self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
        # self.h2.addStretch()
        self.button = QPushButton('Clean', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMinimumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        # self.h2.addStretch()
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
             
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        
        self.effect = QGraphicsOpacityEffect(self)
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 50, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)       
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        # self.mainLayout.setSpacing(30)
        # self.v1.setSpacing(0)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.v1)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)
        self.setLayout(self.mainLayout)
        self.mainLayout.setAlignment(Qt.AlignCenter)
        
        # quit = QAction("Quit", self)
        # quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            if self.var1 is not None:
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.state = self.CheckOne.checkState()
                self.threadpool = QThreadPool()
                self.runner = JobRunnerOne(self.var1,self.state)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Try again')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.var1=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.progress.hide()
        
    def openFileNameDialogOne(self):
        
        fileName, _ = QFileDialog.getOpenFileName(self,"Choose your document ",'',filter="PDF (*.pdf)")
        
        if fileName:        
            if '.pdf' not in fileName:
                fileName=fileName+'.pdf'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxOne.setText(fileName)
            self.var1=self.myTextBoxOne.text()
        return fileName
  
    def alert(self, msg):
        if msg=='Error2':
            self.error('The PDF you are trying to overwrite is currently open. Close it and try again.')
        elif msg=='NoWord':
            self.error('You shall have MS Word installed.')    
        else:
            self.error('Unexpected error found: '+msg)
        self.clean()
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.var1=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('ALL DONE! You can check your document now!')
            self.labelThree.show()
            
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()
     
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None
    
    def instructions(self):
        info = QMessageBox()
        info.setWindowTitle(choices[0][3:])
        
        info.setWindowIcon(QIcon(icon))
        info.setText('''Make sure the PDF/A Compliant check box is selected in MS Word.

If you are not sure, go to: File -> Save as -> PDF -> Options''')

        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setText('OK')
        buttonOk.setFont(fontOne)
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()
        
# <codecell>  
class WorkerSignalsTwo(QObject):
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    
class JobRunnerTwo(QRunnable):    
    signals = WorkerSignalsTwo()
    
    def __init__(self,SaveAs,state,PDFA):
        super().__init__()

        self.is_killed = False 
        self.SaveAs=SaveAs
        self.state=state
        self.PDFA=PDFA
        
    @pyqtSlot()
    def is_opened(self):
        temp_filename=self.SaveAs[:-4]+' temp.pdf'
        if os.path.exists(self.SaveAs) == True:
            try:              
                os.rename(self.SaveAs,temp_filename)
                os.rename(temp_filename,self.SaveAs)               
                return False
            except PermissionError:
                return True
        else:
            return False
        
    def run(self):
        try:
            if self.is_opened() == True:
                self.signals.alert.emit('Error2')
            else:
                backup=shutil.copy(self.SaveAs,self.SaveAs[:-4]+' sehr witzig.pdf')
                src= fitz.open(backup)
                doc = fitz.open()
                for ipage in src:
                    #Some pages, even though having a landscape aspect, are not actually 'landscape', 
                    #just rotated.
                    #These lines of code take care of the many problems that arise when dealing with 
                    #different page sizes. 
                    if ipage.get_contents() != []:               
                        if ipage.rotation==90: 
                            ipage.setRotation(0)               
                        if ipage.rect.width > ipage.rect.height:
                            fmt = fitz.PaperRect("a4-l")  # landscape if input suggests
                        else:
                            fmt = fitz.PaperRect("a4")
                        page = doc.newPage(width = fmt.width, height = fmt.height) 
                        page.showPDFpage(page.rect, src, ipage.number)
                        if page.rect.width > page.rect.height:
                            page.setRotation(90)
                src.close()    
                #-----------------------------------------MINI-LOOP----------------------------------------# 
                done=False
                count=0
                while not done:
                    try:
                        print(count)
                        doc.save(backup,deflate=True)
                        count=5
                        print('sucess!')
                        done=True
                    except RuntimeError:
                        print('Pymupdf permission denied')
                        count+=0.1
                        time.sleep(count)                    
                #-----------------------------------------MINI-LOOP----------------------------------------#        
                doc.close()
                shutil.move(backup,self.SaveAs)
                
                if self.PDFA==2:     
                    done=False
                    count=0
                    while not done:
                        try:
                            backup=shutil.copy(self.SaveAs,self.SaveAs[:-4]+' sehr witzig.pdf')
                            done=True
                        except PermissionError:
                            print('Permission denied')
                            count+=1 
                            time.sleep(count)          
                    time.sleep(1)
                    DestDir=os.path.abspath(os.path.join(self.SaveAs, '..'))                
                    tempDocx='\\tempword.docx'
                    documento = Document()
                    documento.save(DestDir+tempDocx) 
                    try:
                        word = client.DispatchEx('Word.Application')
                    except Exception:
                        self.signals.alert.emit('NoWord')
                    else:
                        worddoc = word.Documents.Open(DestDir+tempDocx,ReadOnly = 1)                   
                        worddoc.SaveAs(self.SaveAs,FileFormat = 17)               
                        worddoc.Close(True)
                        word.Quit()
                        os.remove(DestDir+tempDocx)    
                        pdf=fitz.open(self.SaveAs)
                        opened_file=fitz.open(backup)
                        pdf.insertPDF(opened_file)
                        opened_file.close()
                        pdf.deletePage(0)
                        pdf.saveIncr()
                        pdf.close()               
                        time.sleep(1)      
                        done=False
                        count=0
                        while not done:
                            try:
                                os.remove(backup)
                                done=True
                            except PermissionError:
                                print('Permission denied')
                                count+=1 
                                time.sleep(count)
                if self.state==2:
                    from subprocess import Popen 
                    Popen([self.SaveAs],shell=True)
                time.sleep(1)
                self.signals.finished.emit('Done')
        except Exception as e:      
               self.signals.alert.emit(str(type(e)))                
    def kill(self):
        self.is_killed = True
           
class ActionsTwo(QWidget):
 
    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.initUI()
        self.msg1='Choose a PDF file.'
        
    def initUI(self):
        self.style = QApplication.style()
       
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}")          
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
          
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))
    
        
        self.buttonTwo = QPushButton('Load PDF', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h1.addWidget(self.myTextBoxOne,4)
        
        # self.lineOne = QLabel('/'*250, self) 
        # self.lineOne.setMaximumWidth(800)
        # self.v1.addWidget(self.lineOne)
        
        self.CheckOne = QCheckBox('Open immeddiately', self)  
        self.CheckOne.setFont(fontTwo)
        self.CheckOne.setMinimumHeight(35)
        # self.CheckOne.setMaximumWidth(800)
        self.CheckOne.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}") 
        self.CheckOne.setChecked(True)
        self.v1.addWidget(self.CheckOne)
        
        self.CheckTwo = QCheckBox('Immediately convert to PDF/A', self)
        self.CheckTwo.setFont(fontTwo)
        self.CheckTwo.setMinimumHeight(35)
        # self.CheckTwo.setMaximumWidth(800)
        self.CheckTwo.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}")  
        self.CheckTwo.setChecked(True)
        self.v1.addWidget(self.CheckTwo)
             
        # self.lineTwo = QLabel('/'*250, self)
        # self.lineTwo.setMaximumWidth(800)
        # self.v1.addWidget(self.lineTwo)     
        self.h2.addStretch()
        self.start = QPushButton('Run', self)
        self.start.setStyleSheet(self.style1)
        # self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        # self.start.setMaximumWidth(200)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
    
        self.button = QPushButton('Clean', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMaximumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
        
        
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 100, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)
       
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)
        # self.mainLayout.setSpacing(30)
        # self.v1.setSpacing(0)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.v1)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)       
        self.setLayout(self.mainLayout)
        
        # quit = QAction("Quit", self)
        # quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            if self.var1 is not None:
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.state = self.CheckOne.checkState()
                self.PDFA=self.CheckTwo.checkState()
                self.threadpool = QThreadPool()
                self.runner = JobRunnerTwo(self.var1,self.state,self.PDFA)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Try again')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.var1=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.progress.hide()
        
    def openFileNameDialogOne(self):
    
        fileName, _ = QFileDialog.getOpenFileName(self,"Choose your document",'',filter="PDF (*.pdf)")
        
        if fileName:        
            if '.pdf' not in fileName:
                fileName=fileName+'.pdf'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxOne.setText(fileName)
            self.var1=self.myTextBoxOne.text()
        return fileName
  
    def alert(self, msg):
        if msg=='Error2':
            self.error('The PDF you are trying to overwrite is currently open. Close it and try again.')
        elif msg=='NoWord':
            self.error('You need to have MS Word installed.') 
        else:
            self.error('Unexpected error found: '+msg)
        self.clean()
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.var1=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('ALL DONE! You can check your document now!')
            self.labelThree.show()
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()
    
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None

# <codecell>    
class WorkerSignalsThree(QObject):
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    
class JobRunnerThree(QRunnable):    
    signals = WorkerSignalsThree()
    
    def __init__(self,documents,SaveAs,state,PDFA):
        super().__init__()

        self.is_killed = False 
        self.documents=documents
        self.SaveAs=SaveAs  
        self.state=state
        self.PDFA=PDFA
        
    @pyqtSlot()
    def is_opened(self):
        temp_filename=self.SaveAs[:-4]+' temp.pdf'
        if os.path.exists(self.SaveAs) == True:
            try:              
                os.rename(self.SaveAs,temp_filename)
                os.rename(temp_filename,self.SaveAs)               
                return False
            except PermissionError:
                return True
        else:
            return False
        
    def run(self):
        try:
            if len(self.documents)<2:
                self.signals.alert.emit('Error3')
            elif self.SaveAs in self.documents:
                self.signals.alert.emit('Error1')
            elif self.is_opened() == True:
                self.signals.alert.emit('Error2')
            else:
                pdf=fitz.open()
                for element in self.documents:          
                    opened_file=fitz.open(element)        
                    pdf.insertPDF(opened_file)
                    opened_file.close()          
                pdf.save(self.SaveAs,deflate=True)
                pdf.close()            
                if self.PDFA==2:     
                    done=False
                    count=0
                    while not done:
                        try:
                            backup=shutil.copy(self.SaveAs,self.SaveAs[:-4]+' sehr witzig.pdf')
                            done=True
                        except PermissionError:
                            print('Permission denied')
                            count+=1
                            time.sleep(count)                 
                    time.sleep(1)
                    DestDir=os.path.abspath(os.path.join(self.SaveAs, '..'))                
                    tempDocx='\\tempword.docx'
                    documento = Document()
                    documento.save(DestDir+tempDocx)               
                    word = client.DispatchEx('Word.Application')
                    worddoc = word.Documents.Open(DestDir+tempDocx,ReadOnly = 1)                   
                    worddoc.SaveAs(self.SaveAs,FileFormat = 17)               
                    worddoc.Close(True)
                    word.Quit()
                    os.remove(DestDir+tempDocx)    
                    pdf=fitz.open(self.SaveAs)
                    opened_file=fitz.open(backup)
                    pdf.insertPDF(opened_file)
                    opened_file.close()
                    pdf.deletePage(0)
                    pdf.saveIncr()
                    pdf.close()
                    time.sleep(1)      
                    done=False
                    count=0
                    while not done:
                        try:
                            os.remove(backup)
                            done=True
                        except PermissionError:
                            print('Permission denied')
                            count+=1
                            time.sleep(count)               
                if self.state==2:
                    from subprocess import Popen 
                    Popen([self.SaveAs],shell=True)
                time.sleep(1)
                self.signals.finished.emit('Done')
        except Exception as e:      
               self.signals.alert.emit(str(type(e)))                      
    def kill(self):
        self.is_killed = True
                     
class ActionsThree(QWidget):
 
    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.var2=None
        self.initUI()
        self.msg1='Verify your input.'
        
    def initUI(self):
        self.style = QApplication.style()
        
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}") 
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
          
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.h3=QHBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))
    
       
        self.buttonTwo = QPushButton('Load PDFs', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h1.addWidget(self.myTextBoxOne,4)
        
        self.buttonThree = QPushButton('Save as', self)      
        self.buttonThree.clicked.connect(self.openFileNameDialogTwo)
        self.buttonThree.setMinimumHeight(35)
        # self.buttonThree.setMaximumWidth(200)
        self.buttonThree.setStyleSheet(self.style2)
        self.buttonThree.setFont(fontTwo)
        self.buttonThree.setCursor(QCursor(Qt.PointingHandCursor))
        self.h3.addWidget(self.buttonThree,1)
        
        self.myTextBoxTwo = QLineEdit(self)
        self.myTextBoxTwo.setMinimumHeight(35)  
        self.myTextBoxTwo.setStyleSheet('background-color: rgb(69, 70, 77); color: white')  
        self.myTextBoxTwo.setFont(fontTwo)
        self.myTextBoxTwo.setReadOnly(True)
        self.h3.addWidget(self.myTextBoxTwo,4)
        
        # self.lineOne = QLabel('/'*250, self) 
        # self.lineOne.setMaximumWidth(800)
        # self.v1.addWidget(self.lineOne)
        
        self.CheckOne = QCheckBox('Open immeddiately', self)  
        self.CheckOne.setFont(fontTwo)
        self.CheckOne.setMinimumHeight(35)
        # self.CheckOne.setMaximumWidth(800)
        self.CheckOne.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}") 
        self.CheckOne.setChecked(True)
        self.v1.addWidget(self.CheckOne)
        
        self.CheckTwo = QCheckBox('Immediately convert to PDF/A', self)
        self.CheckTwo.setFont(fontTwo)
        self.CheckTwo.setMinimumHeight(35)
        # self.CheckTwo.setMaximumWidth(800)
        self.CheckTwo.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}")  
        self.CheckTwo.setChecked(False)
        self.v1.addWidget(self.CheckTwo)
             
        # self.lineTwo = QLabel('/'*250, self)
        # self.lineTwo.setMaximumWidth(800)
        # self.v1.addWidget(self.lineTwo)     
        self.h2.addStretch()
        self.start = QPushButton('Run', self)
        self.start.setStyleSheet(self.style1)
        # self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        # self.start.setMaximumWidth(200)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
    
        self.button = QPushButton('Clean', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMaximumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
                
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 50, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)      
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)
        # self.mainLayout.setSpacing(30)
        # self.v1.setSpacing(0)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.h3)
        self.mainLayout.addLayout(self.v1)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)       
        self.setLayout(self.mainLayout)
        
        # quit = QAction("Quit", self)
        # quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            if self.var1 is not None and self.var2 is not None:
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.state = self.CheckOne.checkState()
                self.PDFA=self.CheckTwo.checkState()
                self.threadpool = QThreadPool()
                self.runner = JobRunnerThree(self.var1,self.var2,self.state,self.PDFA)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Try again')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.myTextBoxTwo.setText(None)
        self.var1=None
        self.var2=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.progress.hide()

    def openFileNameDialogOne(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileNames, _ = QFileDialog.getOpenFileNames(self,"Choose your documents","","PDF (*.pdf)",options=options)
        files=[]
        if fileNames: 
            for fileName in fileNames:
                fileName=os.path.abspath(fileName)
                files.append(fileName)
            self.myTextBoxOne.setText(str(files).strip('[').strip(']'))
            self.var1=files
            
        return files
    def openFileNameDialogTwo(self):
        
        fileName, _ = QFileDialog.getSaveFileName(self,"Save as",'',filter="PDF (*.pdf)")
        
        if fileName:        
            if '.pdf' not in fileName:
                fileName=fileName+'.pdf'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxTwo.setText(fileName)
            self.var2=self.myTextBoxTwo.text()
        return fileName
  
    def alert(self, msg):
        if msg=='Error1':
            self.error('Cannot overwrite input documents.')
        elif msg=='Error2':
            self.error('The PDF you are trying to overwrite is currently open. Close it and try again.')
        elif msg=='Error3':
            self.error('Upload two or more PDFs.')
        else:
            self.error('Unexpected error found: '+msg)
        self.clean()
        
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.myTextBoxTwo.setText(None)
            self.var1=None
            self.var2=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('ALL DONE! You can check your document now!')
            
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()
    
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("File Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None
    
    def instructions(self):
        info = QMessageBox()
        info.setWindowTitle(self.title)
        
        info.setWindowIcon(QIcon(icon))
        info.setText("Intrucciones de uso.")
        info.setInformativeText(
        '''
Estas son las instrucciones de uso.'''
        )
        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setText('Entendido')
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()    
    
# <codecell>    
class WorkerSignalsFour(QObject):
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    
class JobRunnerFour(QRunnable):    
    signals = WorkerSignalsFour()
    
    def __init__(self,documents,SaveAs,state):
        super().__init__()

        self.is_killed = False 
        self.documents=documents
        self.SaveAs=SaveAs  
        self.state=state
        
    @pyqtSlot()
    def is_opened(self):
        temp_filename=self.SaveAs[:-4]+' temp.pdf'
        if os.path.exists(self.SaveAs) == True:
            try:              
                os.rename(self.SaveAs,temp_filename)
                os.rename(temp_filename,self.SaveAs)               
                return False
            except PermissionError:
                return True
        else:
            return False
        
    def run(self):
        try:
            if self.SaveAs in self.documents:
                self.signals.alert.emit('Error1')
            elif self.is_opened() == True:
                self.signals.alert.emit('Error2')
            else:
                doc = fitz.open()
                for file in self.documents:
                    img=fitz.open(file)
                    fmt = fitz.PaperRect("a4")
                    pdfbytes = img.convertToPDF()
                    img.close()
                    imgPDF = fitz.open("pdf", pdfbytes)
                    page = doc.newPage(width = fmt.width, height = fmt.height)
                    page.showPDFpage(page.rect, imgPDF, 0)   
                   
                #-----------------------------------------MINI-LOOP----------------------------------------# 
                done=False
                count=0
                while not done:
                    try:
                        print(count)
                        doc.save(self.SaveAs,deflate=True)
                        print('sucess!')
                        done=True
                    except RuntimeError:
                        print('Pymupdf permission denied')
                        count+=0.1
                        time.sleep(count)
                #-----------------------------------------MINI-LOOP----------------------------------------#        
                doc.close()
                if self.state==2:
                    from subprocess import Popen 
                    Popen([self.SaveAs],shell=True)
                time.sleep(1)
                self.signals.finished.emit('Done')
        except Exception as e:      
               self.signals.alert.emit(str(type(e))) 
    def kill(self):
        self.is_killed = True
                    
class ActionsFour(QWidget):

 
    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.var2=None
        self.initUI()
        self.msg1='Verify your inputs.'
        
    def initUI(self):
        self.style = QApplication.style()
       
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}") 
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
          
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.h3=QHBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))
    
      
        self.buttonTwo = QPushButton('Load images', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h1.addWidget(self.myTextBoxOne,4)
        
        self.buttonThree = QPushButton('Save as', self)      
        self.buttonThree.clicked.connect(self.openFileNameDialogTwo)
        self.buttonThree.setMinimumHeight(35)
        # self.buttonThree.setMaximumWidth(200)
        self.buttonThree.setStyleSheet(self.style2)
        self.buttonThree.setFont(fontTwo)
        self.buttonThree.setCursor(QCursor(Qt.PointingHandCursor))
        self.h3.addWidget(self.buttonThree,1)
        
        self.myTextBoxTwo = QLineEdit(self)
        self.myTextBoxTwo.setMinimumHeight(35)  
        self.myTextBoxTwo.setStyleSheet('background-color: rgb(69, 70, 77); color: white')   
        self.myTextBoxTwo.setFont(fontTwo)
        self.myTextBoxTwo.setReadOnly(True)
        self.h3.addWidget(self.myTextBoxTwo,4)
        
        # self.lineOne = QLabel('/'*250, self) 
        # self.lineOne.setMaximumWidth(800)
        # self.v1.addWidget(self.lineOne)
        
        self.CheckOne = QCheckBox('Open immeddiately', self)  
        self.CheckOne.setFont(fontTwo)
        self.CheckOne.setMinimumHeight(35)
        # self.CheckOne.setMaximumWidth(800)
        self.CheckOne.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}") 
        self.CheckOne.setChecked(True)
        self.v1.addWidget(self.CheckOne)
             
        # self.lineTwo = QLabel('/'*250, self)
        # self.lineTwo.setMaximumWidth(800)
        # self.v1.addWidget(self.lineTwo)     
        self.h2.addStretch()
        self.start = QPushButton('Run', self)
        self.start.setStyleSheet(self.style1)
        # self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        # self.start.setMaximumWidth(200)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
    
        self.button = QPushButton('Clean', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMaximumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
        
        
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        
        self.effect = QGraphicsOpacityEffect(self)
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 50, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)       
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)
        # self.mainLayout.setSpacing(30)
        # self.v1.setSpacing(0)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.h3)
        self.mainLayout.addLayout(self.v1)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)       
        self.setLayout(self.mainLayout)
        
        # quit = QAction("Quit", self)
        # quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            if self.var1 is not None and self.var2 is not None:
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.state = self.CheckOne.checkState()
                self.threadpool = QThreadPool()
                self.runner = JobRunnerFour(self.var1,self.var2,self.state)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Try again')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.myTextBoxTwo.setText(None)
        self.var1=None
        self.var2=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.progress.hide()
        
    def openFileNameDialogOne(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileNames, _ = QFileDialog.getOpenFileNames(self,"Choose your images","","Images (*.png *.jpg *.jpeg)",options=options)
        files=[]
        if fileNames: 
            for fileName in fileNames:
                fileName=os.path.abspath(fileName)
                files.append(fileName)
            self.myTextBoxOne.setText(str(files).strip('[').strip(']'))
            self.var1=files
            
        return files
    def openFileNameDialogTwo(self):
    
        fileName, _ = QFileDialog.getSaveFileName(self,"Save as",'',filter="PDF (*.pdf)")
        
        if fileName:        
            if '.pdf' not in fileName:
                fileName=fileName+'.pdf'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxTwo.setText(fileName)
            self.var2=self.myTextBoxTwo.text()
        return fileName
  
    def alert(self, msg):
        if msg=='Error1':
            self.error('Cannot overwrite input documents.')
        elif msg=='Error2':
            self.error('The PDF you are trying to overwrite is currently open. Close it and try again.')
        else:
            self.error('Unexpected error found: '+msg)
        self.clean()
        
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.myTextBoxTwo.setText(None)
            self.var1=None
            self.var2=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('ALL DONE! You can check your document now!')
            self.labelThree.show()
            
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()
    
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None

# <codecell>
class WorkerSignalsFive(QObject):
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    reportMsg=pyqtSignal(str)
 
class JobRunnerFive(QRunnable):    
    signals = WorkerSignalsFive()
    
    def __init__(self,SaveAs,state,choice):
        super().__init__()

        self.is_killed = False 
        self.SaveAs=SaveAs
        self.state=state
        self.choice=choice
        
    @pyqtSlot()
    def is_opened(self):
        temp_filename=self.SaveAs[:-4]+' temp.pdf'
        if os.path.exists(self.SaveAs) == True:
            try:              
                os.rename(self.SaveAs,temp_filename)
                os.rename(temp_filename,self.SaveAs)               
                return False
            except PermissionError:
                return True
        else:
            return False
        
    def run(self):
       
        try:
            if self.is_opened() == True:
                self.signals.alert.emit('Error2')
            else:          
                size_before=os.path.getsize(self.SaveAs)
                size_before/=1000000                           
                import subprocess
                file_name=os.path.basename(self.SaveAs)
                print(file_name)
                output_file=file_name[:-4]+' sehr witzig.pdf'
                os.chdir(os.path.dirname(self.SaveAs))
                quality=''
                if self.choice==0:
                    quality='/printer'
                elif self.choice==1:
                    quality='/ebook'
                elif self.choice==2:
                    quality='/screen'
                startupinfo = None
                if os.name == 'nt':
                    startupinfo = subprocess.STARTUPINFO()
                    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                start=time.time()
                subprocess.run([gs,'-sDEVICE=pdfwrite','-q',
                                        '-dCompatibilityLevel=1.4', 
                                        '-dPDFSETTINGS='+quality,'-dNOSAFER',
                                        '-dNOPAUSE','-dQUIET','-dBATCH',
                                        '-sOutputFile='+'.'+output_file, 
                                        file_name], startupinfo=startupinfo)
                holder=os.path.join(os.path.dirname(self.SaveAs),'.'+output_file)
                if self.is_killed:
                    pass
                else:
                    done=False
                    count=0
                    while not done:
                        try:
                            shutil.copy(holder,self.SaveAs)
                            # shutil.copy(self.SaveAs[:-4]+' sehr witzig.pdf',self.SaveAs)
                            done=True
                        except PermissionError:
                            print('Permission denied')
                            count+=1 
                            time.sleep(count)          
                    time.sleep(1)
                done=False
                count=0
                while not done:
                    try:
                        os.remove(holder)
                        # os.remove(self.SaveAs[:-4]+' sehr witzig.pdf')
                        done=True
                    except PermissionError:
                        print('Permission denied')
                        count+=1 
                        time.sleep(count)
                if self.is_killed:
                    pass
                else:
                    if self.state==2:
                        from subprocess import Popen 
                        Popen([self.SaveAs],shell=True)
                    time.sleep(1)
                    size_after=os.path.getsize(self.SaveAs)
                    size_after/=1000000
                    if (size_before-size_after)<=1:
                        self.signals.reportMsg.emit('Cannot further shrink your document with chosen parameter.')
                    else:
                        self.signals.reportMsg.emit("Your PDF went from %.2f MB to %.2f MB in %.2f sec." % (size_before, size_after,time.time()-start))
                    self.signals.finished.emit('Done')
        except Exception as e:      
            self.signals.alert.emit(str(e))
    def kill(self):
        self.is_killed = True
          
class ActionsFive(QWidget):

    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.initUI()
        self.msg1='Ingresa un archivo PDF.'
        
    def initUI(self):
        
        self.style = QApplication.style()
        
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}")             
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
        self.style4=("QComboBox {selection-background-color: rgb(69, 70, 77);background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);padding-left:10px}"
                     "QComboBox QAbstractItemView::item { min-height: 35px; min-width: 50px;}"
                     "QListView::item { color: white; background-color: rgb(69, 70, 77)}"
                     "QListView::item:selected { color: white; background-color: IndianRed}") 
        self.style5=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}")    

        self.setWindowTitle(self.title)

        self.h4=QHBoxLayout()     
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))

        self.buttonFour = QPushButton('Type', self)  
        self.buttonFour.setMinimumHeight(35)  
        # self.buttonFour.setMaximumWidth(200)
        self.buttonFour.setFont(fontTwo)
        self.buttonFour.setStyleSheet(self.style5)
        self.buttonFour.setEnabled(False)
        self.h4.addWidget(self.buttonFour,1)
        
        self.combo=QComboBox(self)
        self.combo.addItems(['Low',
       'Medium',
       'High'])
        self.combo.setMinimumHeight(35)  
        # self.combo.setMaximumWidth(600)
        self.combo.setFont(fontTwo)
        self.combo.setStyleSheet(self.style4)
        self.listview=QListView()
        self.listview.setFont(fontTwo)
        self.listview.setCursor(QCursor(Qt.PointingHandCursor))
        self.combo.setView(self.listview)
        self.combo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h4.addWidget(self.combo,4)
        
        self.buttonTwo = QPushButton('Load PDF', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        # self.myTextBoxOne.setMaximumWidth(600)
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h1.addWidget(self.myTextBoxOne,4)
        
        # self.lineOne = QLabel('/'*250, self) 
        # self.lineOne.setMaximumWidth(800)
        # self.v1.addWidget(self.lineOne)
        
        self.CheckOne = QCheckBox('Open immediately', self)  
        self.CheckOne.setFont(fontTwo)
        self.CheckOne.setMinimumHeight(35)
        # self.CheckOne.setMaximumWidth(800)
        self.CheckOne.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}") 
        self.CheckOne.setChecked(False)
        self.v1.addWidget(self.CheckOne)
             
        # self.lineTwo = QLabel('/'*250, self)
        # self.lineTwo.setMaximumWidth(800)
        # self.v1.addWidget(self.lineTwo)     
        self.h2.addStretch()
        self.start = QPushButton('Run', self)
        self.start.setStyleSheet(self.style1)
        # self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
        # self.h2.addStretch()
        self.button = QPushButton('Clean', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMinimumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        # self.h2.addStretch()
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
             
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        self.labelOne = QLabel('', self)
        self.labelOne.setFont(fontThree)
        self.labelOne.setAlignment(Qt.AlignCenter)
        self.labelOne.hide()

        self.effect = QGraphicsOpacityEffect(self)
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 50, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)       
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        # self.mainLayout.setSpacing(30)
        # self.v1.setSpacing(0)
        self.mainLayout.addLayout(self.h4)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.v1)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelOne)        
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)
        self.setLayout(self.mainLayout)
        self.mainLayout.setAlignment(Qt.AlignCenter)
        
        # quit = QAction("Quit", self)
        # quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            if self.var1 is not None:
                self.labelTwo.setText('')
                self.labelOne.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.choice = self.combo.currentIndex()
                self.state = self.CheckOne.checkState()
                self.threadpool = QThreadPool()
                self.runner = JobRunnerFive(self.var1,self.state,self.choice)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                    self.runner.signals.reportMsg.disconnect(self.report)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                    self.runner.signals.reportMsg.connect(self.report)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                    self.runner.signals.reportMsg.connect(self.report)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Try again')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.var1=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.progress.hide()
        self.labelOne.setText('')
      
    def openFileNameDialogOne(self):
        
        fileName, _ = QFileDialog.getOpenFileName(self,"Choose your document",'',filter="PDF (*.pdf)")
        
        if fileName:        
            if '.pdf' not in fileName:
                fileName=fileName+'.pdf'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxOne.setText(fileName)
            self.var1=self.myTextBoxOne.text()
        return fileName
  
    def alert(self, msg):
        if msg=='Error2':
            self.error('The PDF you are trying to overwrite is currently open. Close it and try again.')
        else:
            self.error('Unexpected error found: '+msg)
        self.clean()
    def report(self,msg):
        self.labelOne.setText(msg)
        self.labelOne.show()
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.var1=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('ALL DONE! You can check your document now!')
            self.labelThree.show()
            
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()

    # def closeEvent(self, event):
    #     close = QMessageBox()
    #     # close.setWindowTitle(self.title)
    #     close.setWindowTitle("¿Seguro?")
    #     close.setWindowIcon(QIcon(icon))
    #     close.setFont(fontTwo)
    #     close.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
    #     # close.setText("¿Estás seguro?")
    #     # close.setInformativeText('Se detendrá la función si se está ejecutando, pero no te preocupes ya que se guardará el avance.')
    #     close.setText("¿Estás seguro que deseas salir?")
    #     close.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
    #     close = close.exec()

    #     if close == QMessageBox.Yes:           
    #         event.accept()     
    #         self.clean()
    #     else:
    #         event.ignore()
     
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None
    
    def instructions(self):
        info = QMessageBox()
        info.setWindowTitle(choices[9][3:])
        
        info.setWindowIcon(QIcon(icon))
        info.setText('''Low: 300 dpi
Medium: 150 dpi
High: 72 dpi

It can take a couple of minutes to shrink you document, depending on its size and the number of images in it.

As a reference, LuftMensch can shrink 600 MB into a 47 MB document in about 10 minutes.

Do not open your document while is being processed or else all changes will be lost.''')

        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setText('OK')
        buttonOk.setFont(fontOne)
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()

class PicButton(QAbstractButton):
    def __init__(self, pixmap, parent=None):
        super(PicButton, self).__init__(parent)
        self.pixmap = pixmap

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.drawPixmap(event.rect(), self.pixmap)

    def sizeHint(self):
        return self.pixmap.size()             
# <codecell>  
    
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        
        self.window1 = ActionsOne()
        self.window2 = ActionsTwo()
        self.window3 = ActionsThree()
        self.window4 = ActionsFour()
        self.window5 = ActionsFive()
        self.title = 'LuftMensch'
        self.initUI()
        
    def initUI(self):  

        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}"
                     "QPushButton:hover { background-color: rgba(155, 61, 61,230) ;color: white;}"
                      "QPushButton:pressed { background-color: rgb(69, 70, 77) ;color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}"
                      "QPushButton:hover { background-color: rgba(69, 70, 77,230) ;color: white;}"
                      "QPushButton:pressed { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style4=("QComboBox {selection-background-color: rgb(69, 70, 77);background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);padding-left:10px}"
                     "QComboBox QAbstractItemView::item { min-height: 35px; min-width: 50px;}"
                     "QListView::item { color: white; background-color: rgb(69, 70, 77)}"
                     "QListView::item:selected { color: white; background-color: IndianRed}") 
        
        self.style = QApplication.style()
       
        self.setWindowTitle(self.title)       
        # self.setMinimumSize(750,500)
        self.setMinimumSize(530,530)
        # self.resize(500,600)
        # self.move(500, 2)
        # self.setWindowState(Qt.WindowMaximized)
        self.setStyleSheet("background-color: rgb(22, 23, 24); color:CornflowerBlue")
        self.setWindowIcon(QIcon(icon))
        
        self.menuBar = self.menuBar()
        self.menuBar.setCursor(QCursor(Qt.PointingHandCursor))
        self.menuBar.setStyleSheet("QMenuBar {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255)}"
                                   "QMenuBar:item:selected {background-color: white ;color: black}") 
        self.menuBar.addAction('&About', self.about)
        self.menuBar.addAction('&Update', self.update)
        # self.menuBar.addAction('&Ayuda', self.need_help)
        
        self.visitRepo=QMenu("Repository")
        self.visitRepo.setStyleSheet("QMenu {background-color: white; color: black}"
                                   "QMenu:item:selected {background-color: white ;color: rgb(155, 61, 61)}") 
        self.menuBar.addMenu(self.visitRepo)
        self.visitRepo.setCursor(QCursor(Qt.PointingHandCursor))
        self.visitRepo.addAction('&Visit repository', self.repo)
        
        self.help=QMenu("&Instructions")
        self.help.setStyleSheet("QMenu {background-color: white; color: black}"
                                   "QMenu:item:selected {background-color: white ;color: rgb(155, 61, 61)}") 
        self.menuBar.addMenu(self.help)
        self.help.setCursor(QCursor(Qt.PointingHandCursor))
        self.help.addAction(choices[0], self.window1.instructions)
        self.help.addAction(choices[4], self.window5.instructions)
   
        self.stackedLayout = QStackedLayout()
              
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)    
    
        self.h=QHBoxLayout()
        self.v=QVBoxLayout()
        
        self.v0=QVBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
        self.v3=QVBoxLayout()
        self.h3=QHBoxLayout()
        self.h4=QHBoxLayout()       
        self.h5=QHBoxLayout()
        self.h10=QVBoxLayout()

        windows=[self.window1,
                 self.window2,
                 self.window3,
                 self.window4,
                 self.window5]    
        
        for window in windows:
            self.stackedLayout.addWidget(window)
            
        self.pageCombo = QComboBox()   
        self.pageCombo.addItems(choices)
        self.pageCombo.setMinimumHeight(35)
        self.pageCombo.setStyleSheet(self.style4)
        self.listview=QListView()
        self.listview.setFont(fontTwo)
        self.listview.setCursor(QCursor(Qt.PointingHandCursor))
        self.pageCombo.setView(self.listview)
        self.pageCombo.setCursor(QCursor(Qt.PointingHandCursor))
        self.pageCombo.setFont(fontTwo)
        self.pageCombo.activated.connect(self.toggle_window)

        self.v0.addWidget(self.pageCombo)

        self.h.addLayout(self.v1)
        self.h.addLayout(self.v)
        self.h.addLayout(self.v2)   

        self.stackedLayout.setAlignment(Qt.AlignCenter)
        self.h.setAlignment(Qt.AlignCenter)
               
        self.mainLayout.addLayout(self.h,1)   
        self.mainLayout.addLayout(self.v0,0)   
        
        self.mainLayout.addLayout(self.stackedLayout,6)        
      
        self.pixmap = QPixmap(icon)
        self.pixmap = self.pixmap.scaled(70, 70, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setPixmap(self.pixmap) 
        self.labelThree.setAlignment(Qt.AlignCenter) 
        self.v1.addWidget(self.labelThree)
        
        self.logo = QPixmap(logo)
        self.logo = self.logo.scaled(110, 110, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelFour = QLabel('', self)
        self.labelFour.setPixmap(self.logo) 
        self.labelFour.setAlignment(Qt.AlignCenter) 
        self.v.addWidget(self.labelFour)
        
        self.titleOne = QLabel('Version 1.4.3', self)
        self.titleOne.setFont(fontFive)
        self.titleOne.setStyleSheet("color:	IndianRed")
        self.titleOne.setAlignment(Qt.AlignRight | Qt.AlignBottom)  
        self.v2.addWidget(self.titleOne)
        
        self.labelOne = QLabel('Hi there, '+username, self)
        self.labelOne.setFont(fontFive)
        self.labelOne.setAlignment(Qt.AlignRight)  
        self.v2.addWidget(self.labelOne)        
        
        self.status_label = QLabel()
        self.statusBar().addPermanentWidget(self.status_label)
        self.status_label.setText('Version 1.4.3 released in September 2021')
        
        self.btn = PicButton(QPixmap(donate_pic).scaled(150,50))
        self.btn.setCursor(QCursor(Qt.PointingHandCursor))
        self.btn.clicked.connect(self.donate)
        self.h10.addWidget(self.btn)  
        self.h10.addStretch(0)  
        self.h10.setAlignment(Qt.AlignCenter)
        self.mainLayout.addLayout(self.h10)

        self.w = QWidget(self)
        self.w.setLayout(self.mainLayout)
        self.setCentralWidget(self.w)
        
        quit = QAction("Quit", self)
        quit.triggered.connect(self.closeEvent)
    def donate(self):
        op('https://www.paypal.com/donate?token=Au3wVcTk2LBobTkSfVdJmlH15ru5nygT6tVW_0o-PoCn-klgwrE6sfVuzTtMKn5dMHbDEXdUSEZq8ypN')
 
    def toggle_window(self):
        self.stackedLayout.setCurrentIndex(self.pageCombo.currentIndex())
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()

    def closeEvent(self, event):
        close = QMessageBox()
        close.setWindowTitle("Are you sure?")
        close.setWindowIcon(QIcon(icon))
        close.setFont(fontTwo)
        close.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        if self.window5.runner: 
            close.setText("LuftMensch is shrinking your document. If you quit now, all progress will be lost.")
        else:          
            close.setText("You are about to quit LuftMensch.")           
        close.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        buttonYes = close.button(QMessageBox.Yes)
        buttonYes.setCursor(QCursor(Qt.PointingHandCursor))
        buttonYes.setFont(fontOne)
        buttonYes.setText('Yes')
        buttonCancel = close.button(QMessageBox.Cancel)
        buttonCancel.setText('No')
        buttonCancel.setCursor(QCursor(Qt.PointingHandCursor))
        buttonCancel.setFont(fontOne)
        close = close.exec()

        if close == QMessageBox.Yes:  
            if self.window5.runner: 
                self.window5.runner.kill()                
            event.accept() 
        else:
            event.ignore()
    def repo(self):
        op('https://github.com/lheredias/Luftmensch')
    def about(self):
        info = QMessageBox()
        info.setWindowTitle("About LuftMensch")
        
        info.setWindowIcon(QIcon(icon))
        info.setText('''LuftMensch is a free, open source application aimed at dealing with common PDF-related tasks not easily available without a license purchase or paid subscription.''')

        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setText('Understood')
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()    
     
    def update(self):
        info = QMessageBox()
        info.setWindowTitle("¿How to update LuftMensch?")
        
        info.setWindowIcon(QIcon(icon))
        info.setText('''Download the latest version (in English) and overwrite the current one.''')

        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        buttonYes = info.button(QMessageBox.Yes)
        buttonYes.setCursor(QCursor(Qt.PointingHandCursor))
        buttonYes.setText('Releases')
        buttonYes.setFont(fontOne)
        buttonCancel = info.button(QMessageBox.Cancel)
        buttonCancel.setCursor(QCursor(Qt.PointingHandCursor))
        buttonCancel.setText('Understood')
        buttonCancel.setFont(fontOne)
        info.setDefaultButton(QMessageBox.Cancel)
        info.show()
        retval = info.exec_()
        print(retval)
        if retval==16384:
            op('https://github.com/lheredias/Luftmensch/releases')
    
if __name__ == '__main__':
    os.environ["QT_AUTO_SCREEN_SCALE_FACTOR"] = "1"
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    app.setAttribute(Qt.AA_EnableHighDpiScaling,True)
    app.setAttribute(Qt.AA_UseHighDpiPixmaps, True)
    app.setWindowIcon(QIcon(icon))
    w = MainWindow()
    w.show() 
    sys.exit(app.exec_())
    
