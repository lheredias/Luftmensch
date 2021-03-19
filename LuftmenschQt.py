from PyQt5.QtCore import (QThread, pyqtSignal,QThreadPool,pyqtSlot,QRunnable,QObject,Qt)
from PyQt5.QtWidgets import (QApplication, QDialog, QMainWindow,QLabel,QFileDialog,QAction,
                             QProgressBar, QPushButton,QMessageBox,QLineEdit,QMenu,
                             QErrorMessage,QGraphicsOpacityEffect,QComboBox,QHBoxLayout,
                             QGridLayout,QSpacerItem,QSizePolicy,QFrame,QStackedLayout,
                             QTextEdit,QTableView,QStyle,QCheckBox,QVBoxLayout,QWidget,
                             QDesktopWidget,QListView,QStyleFactory)

from PyQt5.QtGui import (QIcon,QFont,QPixmap,QCursor)
import time
import sys
import os
import getpass
import shutil
import fitz
from win32com import client
from docx import Document  
from re import findall
# from selenium import webdriver
# from selenium.webdriver.support import expected_conditions as ec
# from selenium.webdriver.common.by import By
# from selenium.webdriver.support.ui import WebDriverWait
# from selenium.webdriver.support.select import Select
# from selenium.common.exceptions import TimeoutException
# from selenium.common.exceptions import WebDriverException
# import pandas as pd
# import openpyxl
# from PIL import Image, ImageOps 
if hasattr(Qt, 'AA_EnableHighDpiScaling'):
    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)

if hasattr(Qt, 'AA_UseHighDpiPixmaps'):
    QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    
    return os.path.join(base_path, relative_path)

icon=resource_path('finalicon.ico')
pic=resource_path('check small.png')
logo=resource_path('app name.png')
username=getpass.getuser()
defaultDir='D:\\Usuarios\\'+username+'\\Documents'

choices=['1. Convertir PDF en PDF/A',
         '2. Obtener PDF con páginas del mismo tamaño',
                 '3. Unir varios archivos PDF',
                 '4. Convertir una o varias imágenes en un solo archivo PDF',
                 '5. Crear archivo .zip de Requerimientos y Cartas',
                 '6. Crear archivo .zip de Valores',
                 '7. Generar archivo de texto para solicitar descarga de LE'] 
fontOne = QFont("Helvetica", 9)
fontTwo=QFont("Helvetica", 9)
fontThree=QFont('Consolas', 11)  #Done message font
fontFive=QFont('Consolas', 11) #Version font
# <codecell>
class WorkerSignalsOne(QObject):
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    
class JobRunnerOne(QRunnable):    
    signals = WorkerSignalsOne()
    
    def __init__(self,SaveAs,state):
        super().__init__()

        self.is_killed = False 
        self.SaveAs=SaveAs
        self.state=state
        
    @pyqtSlot()
    def is_opened(self):
        temp_filename=self.SaveAs[:-4]+' temp.pdf'
        if os.path.exists(self.SaveAs) == True:
            try:              
                os.rename(self.SaveAs,temp_filename)
                os.rename(temp_filename,self.SaveAs)               
                return False
            except PermissionError:
                return True
        else:
            return False
        
    def run(self):
       
        
        if self.is_opened() == True:
            self.signals.alert.emit('Error2')
        else:
            
                
            done=False
            count=0
            while not done:
                try:
                    backup=shutil.copy(self.SaveAs,self.SaveAs[:-4]+' sehr witzig.pdf')
                    done=True
                except PermissionError:
                    print('Permission denied')
                    count+=1 
                    time.sleep(count)          
            time.sleep(1)
            DestDir=os.path.abspath(os.path.join(self.SaveAs, '..'))                
            tempDocx='\\tempword.docx'
            documento = Document()
            documento.save(DestDir+tempDocx)               
            word = client.DispatchEx('Word.Application')
            worddoc = word.Documents.Open(DestDir+tempDocx,ReadOnly = 1)                   
            worddoc.SaveAs(self.SaveAs,FileFormat = 17)               
            worddoc.Close(True)
            word.Quit()
            os.remove(DestDir+tempDocx)    
            pdf=fitz.open(self.SaveAs)
            opened_file=fitz.open(backup)
            pdf.insertPDF(opened_file)
            opened_file.close()
            pdf.deletePage(0)
            pdf.saveIncr()
            pdf.close()               
            time.sleep(1)      
            done=False
            count=0
            while not done:
                try:
                    os.remove(backup)
                    done=True
                except PermissionError:
                    print('Permission denied')
                    count+=1 
                    time.sleep(count)
            if self.state==2:
                from subprocess import Popen 
                Popen([self.SaveAs],shell=True)
            time.sleep(1)
            self.signals.finished.emit('Done')
                            
    def kill(self):
        self.is_killed = True
          
class ActionsOne(QWidget):

    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.initUI()
        self.msg1='Ingresa un archivo PDF.'
        
    def initUI(self):
        
        self.style = QApplication.style()
        
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}")             
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
          
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))

        
        self.buttonTwo = QPushButton('Cargar PDF', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        # self.myTextBoxOne.setMaximumWidth(600)
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h1.addWidget(self.myTextBoxOne,4)
        
        # self.lineOne = QLabel('/'*250, self) 
        # self.lineOne.setMaximumWidth(800)
        # self.v1.addWidget(self.lineOne)
        
        self.CheckOne = QCheckBox('Abrir de inmediato el documento generado', self)  
        self.CheckOne.setFont(fontTwo)
        self.CheckOne.setMinimumHeight(35)
        # self.CheckOne.setMaximumWidth(800)
        self.CheckOne.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}") 
        self.CheckOne.setChecked(True)
        self.v1.addWidget(self.CheckOne)
             
        # self.lineTwo = QLabel('/'*250, self)
        # self.lineTwo.setMaximumWidth(800)
        # self.v1.addWidget(self.lineTwo)     
        self.h2.addStretch()
        self.start = QPushButton('Ejecutar', self)
        self.start.setStyleSheet(self.style1)
        self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
        # self.h2.addStretch()
        self.button = QPushButton('Limpiar', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMinimumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        # self.h2.addStretch()
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
             
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        
        self.effect = QGraphicsOpacityEffect(self)
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 50, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)       
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        # self.mainLayout.setSpacing(30)
        # self.v1.setSpacing(0)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.v1)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)
        self.setLayout(self.mainLayout)
        self.mainLayout.setAlignment(Qt.AlignCenter)
        
        quit = QAction("Quit", self)
        quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            if self.var1 is not None:
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.state = self.CheckOne.checkState()
                self.threadpool = QThreadPool()
                self.runner = JobRunnerOne(self.var1,self.state)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Intenta de nuevo.')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.var1=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.progress.hide()
        
    def openFileNameDialogOne(self):
        
        fileName, _ = QFileDialog.getOpenFileName(self,"Selecciona tu documento",'',filter="PDF (*.pdf)")
        
        if fileName:        
            if '.pdf' not in fileName:
                fileName=fileName+'.pdf'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxOne.setText(fileName)
            self.var1=self.myTextBoxOne.text()
        return fileName
  
    def alert(self, msg):
        if msg=='Error2':
            self.error('Cierra el PDF sobre el cual deseas guardar el resultado.')
    
        self.clean()
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.var1=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('¡Listo, ya puedes visualizar tus documentos!')
            self.labelThree.show()
            
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()

    def closeEvent(self, event):
        close = QMessageBox()
        # close.setWindowTitle(self.title)
        close.setWindowTitle("¿Seguro?")
        close.setWindowIcon(QIcon(icon))
        close.setFont(fontTwo)
        close.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        # close.setText("¿Estás seguro?")
        # close.setInformativeText('Se detendrá la función si se está ejecutando, pero no te preocupes ya que se guardará el avance.')
        close.setText("¿Estás seguro que deseas salir?")
        close.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        close = close.exec()

        if close == QMessageBox.Yes:           
            event.accept()     
            self.clean()
        else:
            event.ignore()
     
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None
    
    def instructions(self):
        info = QMessageBox()
        info.setWindowTitle(choices[0][3:])
        
        info.setWindowIcon(QIcon(icon))
        info.setText('''La opción "Compatible con PDF/A" debe encontrarse activa en Microsoft Word. Para activarla, dirígete a:

Archivo -> Guardar como -> PDF -> Opciones''')

        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setText('Entendido')
        buttonOk.setFont(fontOne)
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()
        
# <codecell>  
class WorkerSignalsTwo(QObject):
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    
class JobRunnerTwo(QRunnable):    
    signals = WorkerSignalsTwo()
    
    def __init__(self,SaveAs,state,PDFA):
        super().__init__()

        self.is_killed = False 
        self.SaveAs=SaveAs
        self.state=state
        self.PDFA=PDFA
        
    @pyqtSlot()
    def is_opened(self):
        temp_filename=self.SaveAs[:-4]+' temp.pdf'
        if os.path.exists(self.SaveAs) == True:
            try:              
                os.rename(self.SaveAs,temp_filename)
                os.rename(temp_filename,self.SaveAs)               
                return False
            except PermissionError:
                return True
        else:
            return False
        
    def run(self):
        if self.is_opened() == True:
            self.signals.alert.emit('Error2')
        else:
            backup=shutil.copy(self.SaveAs,self.SaveAs[:-4]+' sehr witzig.pdf')
            src= fitz.open(backup)
            doc = fitz.open()
            for ipage in src:
                #Some pages, even though having a landscape aspect, are not actually 'landscape', 
                #just rotated.
                #These lines of code take care of the many problems that arise when dealing with 
                #different page sizes. 
                if ipage.get_contents() != []:               
                    if ipage.rotation==90: 
                        ipage.setRotation(0)               
                    if ipage.rect.width > ipage.rect.height:
                        fmt = fitz.PaperRect("a4-l")  # landscape if input suggests
                    else:
                        fmt = fitz.PaperRect("a4")
                    page = doc.newPage(width = fmt.width, height = fmt.height) 
                    page.showPDFpage(page.rect, src, ipage.number)
                    if page.rect.width > page.rect.height:
                        page.setRotation(90)
            src.close()    
            #-----------------------------------------MINI-LOOP----------------------------------------# 
            done=False
            count=0
            while not done:
                try:
                    print(count)
                    doc.save(backup,deflate=True)
                    count=5
                    print('sucess!')
                    done=True
                except RuntimeError:
                    print('Pymupdf permission denied')
                    count+=0.1
                    time.sleep(count)                    
            #-----------------------------------------MINI-LOOP----------------------------------------#        
            doc.close()
            shutil.move(backup,self.SaveAs)
            
            if self.PDFA==2:
                from win32com import client
                from docx import Document      
                done=False
                count=0
                while not done:
                    try:
                        backup=shutil.copy(self.SaveAs,self.SaveAs[:-4]+' sehr witzig.pdf')
                        done=True
                    except PermissionError:
                        print('Permission denied')
                        count+=1 
                        time.sleep(count)          
                time.sleep(1)
                DestDir=os.path.abspath(os.path.join(self.SaveAs, '..'))                
                tempDocx='\\tempword.docx'
                documento = Document()
                documento.save(DestDir+tempDocx)               
                word = client.DispatchEx('Word.Application')
                worddoc = word.Documents.Open(DestDir+tempDocx,ReadOnly = 1)                   
                worddoc.SaveAs(self.SaveAs,FileFormat = 17)               
                worddoc.Close(True)
                word.Quit()
                os.remove(DestDir+tempDocx)    
                pdf=fitz.open(self.SaveAs)
                opened_file=fitz.open(backup)
                pdf.insertPDF(opened_file)
                opened_file.close()
                pdf.deletePage(0)
                pdf.saveIncr()
                pdf.close()               
                time.sleep(1)      
                done=False
                count=0
                while not done:
                    try:
                        os.remove(backup)
                        done=True
                    except PermissionError:
                        print('Permission denied')
                        count+=1 
                        time.sleep(count)
            if self.state==2:
                from subprocess import Popen 
                Popen([self.SaveAs],shell=True)
            time.sleep(1)
            self.signals.finished.emit('Done')
                            
    def kill(self):
        self.is_killed = True
           
class ActionsTwo(QWidget):
 
    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.initUI()
        self.msg1='Ingresa un archivo PDF.'
        
    def initUI(self):
        self.style = QApplication.style()
       
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}")          
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
          
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))
    
        
        self.buttonTwo = QPushButton('Cargar PDF', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h1.addWidget(self.myTextBoxOne,4)
        
        # self.lineOne = QLabel('/'*250, self) 
        # self.lineOne.setMaximumWidth(800)
        # self.v1.addWidget(self.lineOne)
        
        self.CheckOne = QCheckBox('Abrir de inmediato el documento generado', self)  
        self.CheckOne.setFont(fontTwo)
        self.CheckOne.setMinimumHeight(35)
        # self.CheckOne.setMaximumWidth(800)
        self.CheckOne.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}") 
        self.CheckOne.setChecked(True)
        self.v1.addWidget(self.CheckOne)
        
        self.CheckTwo = QCheckBox('Convertir de inmediato a PDF/A', self)
        self.CheckTwo.setFont(fontTwo)
        self.CheckTwo.setMinimumHeight(35)
        # self.CheckTwo.setMaximumWidth(800)
        self.CheckTwo.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}")  
        self.CheckTwo.setChecked(True)
        self.v1.addWidget(self.CheckTwo)
             
        # self.lineTwo = QLabel('/'*250, self)
        # self.lineTwo.setMaximumWidth(800)
        # self.v1.addWidget(self.lineTwo)     
        self.h2.addStretch()
        self.start = QPushButton('Ejecutar', self)
        self.start.setStyleSheet(self.style1)
        self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        # self.start.setMaximumWidth(200)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
    
        self.button = QPushButton('Limpiar', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMaximumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
        
        
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 100, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)
       
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)
        # self.mainLayout.setSpacing(30)
        # self.v1.setSpacing(0)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.v1)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)       
        self.setLayout(self.mainLayout)
        
        quit = QAction("Quit", self)
        quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            if self.var1 is not None:
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.state = self.CheckOne.checkState()
                self.PDFA=self.CheckTwo.checkState()
                self.threadpool = QThreadPool()
                self.runner = JobRunnerTwo(self.var1,self.state,self.PDFA)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Intenta de nuevo.')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.var1=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.progress.hide()
        
    def openFileNameDialogOne(self):
    
        fileName, _ = QFileDialog.getOpenFileName(self,"Selecciona tu documento",'',filter="PDF (*.pdf)")
        
        if fileName:        
            if '.pdf' not in fileName:
                fileName=fileName+'.pdf'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxOne.setText(fileName)
            self.var1=self.myTextBoxOne.text()
        return fileName
  
    def alert(self, msg):
        if msg=='Error2':
            self.error('Cierra el PDF sobre el cual intentas guardar el resultado.')
        self.clean()
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.var1=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('¡Listo, ya puedes visualizar tus documentos!')
            self.labelThree.show()
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()

    def closeEvent(self, event):
        close = QMessageBox()
        # close.setWindowTitle(self.title)
        close.setWindowTitle("¿Seguro?")
        close.setWindowIcon(QIcon(icon))
        close.setFont(fontTwo)
        close.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        # close.setText("¿Estás seguro?")
        # close.setInformativeText('Se detendrá la función si se está ejecutando, pero no te preocupes ya que se guardará el avance.')
        close.setText("¿Estás seguro que deseas salir?")
        close.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        close = close.exec()

        if close == QMessageBox.Yes:           
            event.accept()            
            self.clean()
        else:
            event.ignore()  
    
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None
    
    def instructions(self):
        info = QMessageBox()
        info.setWindowTitle(self.title)
        
        info.setWindowIcon(QIcon(icon))
        info.setText("Intrucciones de uso.")
        info.setInformativeText(
        '''
Estas son las instrucciones de uso.'''
        )
        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setText('Entendido')
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()
   
# <codecell>    
class WorkerSignalsThree(QObject):
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    
class JobRunnerThree(QRunnable):    
    signals = WorkerSignalsThree()
    
    def __init__(self,documents,SaveAs,state,PDFA):
        super().__init__()

        self.is_killed = False 
        self.documents=documents
        self.SaveAs=SaveAs  
        self.state=state
        self.PDFA=PDFA
        
    @pyqtSlot()
    def is_opened(self):
        temp_filename=self.SaveAs[:-4]+' temp.pdf'
        if os.path.exists(self.SaveAs) == True:
            try:              
                os.rename(self.SaveAs,temp_filename)
                os.rename(temp_filename,self.SaveAs)               
                return False
            except PermissionError:
                return True
        else:
            return False
        
    def run(self):
        if len(self.documents)<2:
            self.signals.alert.emit('Error3')
        elif self.SaveAs in self.documents:
            self.signals.alert.emit('Error1')
        elif self.is_opened() == True:
            self.signals.alert.emit('Error2')
        else:
            pdf=fitz.open()
            for element in self.documents:          
                opened_file=fitz.open(element)        
                pdf.insertPDF(opened_file)
                opened_file.close()          
            pdf.save(self.SaveAs,deflate=True)
            pdf.close()            
            if self.PDFA==2:
                from win32com import client
                from docx import Document        
                done=False
                count=0
                while not done:
                    try:
                        backup=shutil.copy(self.SaveAs,self.SaveAs[:-4]+' sehr witzig.pdf')
                        done=True
                    except PermissionError:
                        print('Permission denied')
                        count+=1
                        time.sleep(count)                 
                time.sleep(1)
                DestDir=os.path.abspath(os.path.join(self.SaveAs, '..'))                
                tempDocx='\\tempword.docx'
                documento = Document()
                documento.save(DestDir+tempDocx)               
                word = client.DispatchEx('Word.Application')
                worddoc = word.Documents.Open(DestDir+tempDocx,ReadOnly = 1)                   
                worddoc.SaveAs(self.SaveAs,FileFormat = 17)               
                worddoc.Close(True)
                word.Quit()
                os.remove(DestDir+tempDocx)    
                pdf=fitz.open(self.SaveAs)
                opened_file=fitz.open(backup)
                pdf.insertPDF(opened_file)
                opened_file.close()
                pdf.deletePage(0)
                pdf.saveIncr()
                pdf.close()
                time.sleep(1)      
                done=False
                count=0
                while not done:
                    try:
                        os.remove(backup)
                        done=True
                    except PermissionError:
                        print('Permission denied')
                        count+=1
                        time.sleep(count)               
            if self.state==2:
                from subprocess import Popen 
                Popen([self.SaveAs],shell=True)
            time.sleep(1)
            self.signals.finished.emit('Done')
                            
    def kill(self):
        self.is_killed = True
                     
class ActionsThree(QWidget):
 
    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.var2=None
        self.initUI()
        self.msg1='Verifica los datos ingresados.'
        
    def initUI(self):
        self.style = QApplication.style()
        
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}") 
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
          
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.h3=QHBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))
    
       
        self.buttonTwo = QPushButton('Cargar PDF', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h1.addWidget(self.myTextBoxOne,4)
        
        self.buttonThree = QPushButton('Guardar como', self)      
        self.buttonThree.clicked.connect(self.openFileNameDialogTwo)
        self.buttonThree.setMinimumHeight(35)
        # self.buttonThree.setMaximumWidth(200)
        self.buttonThree.setStyleSheet(self.style2)
        self.buttonThree.setFont(fontTwo)
        self.buttonThree.setCursor(QCursor(Qt.PointingHandCursor))
        self.h3.addWidget(self.buttonThree,1)
        
        self.myTextBoxTwo = QLineEdit(self)
        self.myTextBoxTwo.setMinimumHeight(35)  
        self.myTextBoxTwo.setStyleSheet('background-color: rgb(69, 70, 77); color: white')  
        self.myTextBoxTwo.setFont(fontTwo)
        self.myTextBoxTwo.setReadOnly(True)
        self.h3.addWidget(self.myTextBoxTwo,4)
        
        # self.lineOne = QLabel('/'*250, self) 
        # self.lineOne.setMaximumWidth(800)
        # self.v1.addWidget(self.lineOne)
        
        self.CheckOne = QCheckBox('Abrir de inmediato el documento generado', self)  
        self.CheckOne.setFont(fontTwo)
        self.CheckOne.setMinimumHeight(35)
        # self.CheckOne.setMaximumWidth(800)
        self.CheckOne.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}") 
        self.CheckOne.setChecked(True)
        self.v1.addWidget(self.CheckOne)
        
        self.CheckTwo = QCheckBox('Convertir de inmediato a PDF/A', self)
        self.CheckTwo.setFont(fontTwo)
        self.CheckTwo.setMinimumHeight(35)
        # self.CheckTwo.setMaximumWidth(800)
        self.CheckTwo.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}")  
        self.CheckTwo.setChecked(False)
        self.v1.addWidget(self.CheckTwo)
             
        # self.lineTwo = QLabel('/'*250, self)
        # self.lineTwo.setMaximumWidth(800)
        # self.v1.addWidget(self.lineTwo)     
        self.h2.addStretch()
        self.start = QPushButton('Ejecutar', self)
        self.start.setStyleSheet(self.style1)
        self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        # self.start.setMaximumWidth(200)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
    
        self.button = QPushButton('Limpiar', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMaximumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
                
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 50, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)      
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)
        # self.mainLayout.setSpacing(30)
        # self.v1.setSpacing(0)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.h3)
        self.mainLayout.addLayout(self.v1)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)       
        self.setLayout(self.mainLayout)
        
        quit = QAction("Quit", self)
        quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            if self.var1 is not None and self.var2 is not None:
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.state = self.CheckOne.checkState()
                self.PDFA=self.CheckTwo.checkState()
                self.threadpool = QThreadPool()
                self.runner = JobRunnerThree(self.var1,self.var2,self.state,self.PDFA)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Intenta de nuevo.')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.myTextBoxTwo.setText(None)
        self.var1=None
        self.var2=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.progress.hide()

    def openFileNameDialogOne(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileNames, _ = QFileDialog.getOpenFileNames(self,"Selecciona tus PDFs","","PDF (*.pdf)",options=options)
        files=[]
        if fileNames: 
            for fileName in fileNames:
                fileName=os.path.abspath(fileName)
                files.append(fileName)
            self.myTextBoxOne.setText(str(files).strip('[').strip(']'))
            self.var1=files
            
        return files
    def openFileNameDialogTwo(self):
        
        fileName, _ = QFileDialog.getSaveFileName(self,"Guardar como",'',filter="PDF (*.pdf)")
        
        if fileName:        
            if '.pdf' not in fileName:
                fileName=fileName+'.pdf'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxTwo.setText(fileName)
            self.var2=self.myTextBoxTwo.text()
        return fileName
  
    def alert(self, msg):
        if msg=='Error1':
            self.error('No puedes guardar el resultado encima de uno de los PDFs que deseas unir.')
        elif msg=='Error2':
            self.error('Cierra el PDF sobre el cual deseas guardar el resultado.')
        elif msg=='Error3':
            self.error('Debes seleccionar más de un PDF.')
        self.clean()
        
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.myTextBoxTwo.setText(None)
            self.var1=None
            self.var2=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('¡Listo, ya puedes visualizar tus documentos!')
            
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()

    def closeEvent(self, event):
        close = QMessageBox()
        # close.setWindowTitle(self.title)
        close.setWindowTitle("¿Seguro?")
        close.setWindowIcon(QIcon(icon))
        close.setFont(fontTwo)
        close.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        # close.setText("¿Estás seguro?")
        # close.setInformativeText('Se detendrá la función si se está ejecutando, pero no te preocupes ya que se guardará el avance.')
        close.setText("¿Estás seguro que deseas salir?")
        close.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        close = close.exec()

        if close == QMessageBox.Yes:           
            event.accept()         
            self.clean()
        else:
            event.ignore() 
    
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("File Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None
    
    def instructions(self):
        info = QMessageBox()
        info.setWindowTitle(self.title)
        
        info.setWindowIcon(QIcon(icon))
        info.setText("Intrucciones de uso.")
        info.setInformativeText(
        '''
Estas son las instrucciones de uso.'''
        )
        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setText('Entendido')
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()    
    
# <codecell>    
class WorkerSignalsFour(QObject):
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    
class JobRunnerFour(QRunnable):    
    signals = WorkerSignalsFour()
    
    def __init__(self,documents,SaveAs,state):
        super().__init__()

        self.is_killed = False 
        self.documents=documents
        self.SaveAs=SaveAs  
        self.state=state
        
    @pyqtSlot()
    def is_opened(self):
        temp_filename=self.SaveAs[:-4]+' temp.pdf'
        if os.path.exists(self.SaveAs) == True:
            try:              
                os.rename(self.SaveAs,temp_filename)
                os.rename(temp_filename,self.SaveAs)               
                return False
            except PermissionError:
                return True
        else:
            return False
        
    def run(self):
    
        if self.SaveAs in self.documents:
            self.signals.alert.emit('Error1')
        elif self.is_opened() == True:
            self.signals.alert.emit('Error2')
        else:
            doc = fitz.open()
            for file in self.documents:
                img=fitz.open(file)
                fmt = fitz.PaperRect("a4")
                pdfbytes = img.convertToPDF()
                img.close()
                imgPDF = fitz.open("pdf", pdfbytes)
                page = doc.newPage(width = fmt.width, height = fmt.height)
                page.showPDFpage(page.rect, imgPDF, 0)   
               
            #-----------------------------------------MINI-LOOP----------------------------------------# 
            done=False
            count=0
            while not done:
                try:
                    print(count)
                    doc.save(self.SaveAs,deflate=True)
                    print('sucess!')
                    done=True
                except RuntimeError:
                    print('Pymupdf permission denied')
                    count+=0.1
                    time.sleep(count)
            #-----------------------------------------MINI-LOOP----------------------------------------#        
            doc.close()
            if self.state==2:
                from subprocess import Popen 
                Popen([self.SaveAs],shell=True)
            time.sleep(1)
            self.signals.finished.emit('Done')
            
    def kill(self):
        self.is_killed = True
                    
class ActionsFour(QWidget):
    """
    Simple dialog that consists of a Progress Bar and a Button.
    Clicking on the button results in the start of a timer and
    updates the progress bar.
    """
 
    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.var2=None
        self.initUI()
        self.msg1='Verifica los datos ingresados.'
        
    def initUI(self):
        self.style = QApplication.style()
       
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}") 
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
          
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.h3=QHBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))
    
      
        self.buttonTwo = QPushButton('Cargar imágenes', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h1.addWidget(self.myTextBoxOne,4)
        
        self.buttonThree = QPushButton('Guardar como', self)      
        self.buttonThree.clicked.connect(self.openFileNameDialogTwo)
        self.buttonThree.setMinimumHeight(35)
        # self.buttonThree.setMaximumWidth(200)
        self.buttonThree.setStyleSheet(self.style2)
        self.buttonThree.setFont(fontTwo)
        self.buttonThree.setCursor(QCursor(Qt.PointingHandCursor))
        self.h3.addWidget(self.buttonThree,1)
        
        self.myTextBoxTwo = QLineEdit(self)
        self.myTextBoxTwo.setMinimumHeight(35)  
        self.myTextBoxTwo.setStyleSheet('background-color: rgb(69, 70, 77); color: white')   
        self.myTextBoxTwo.setFont(fontTwo)
        self.myTextBoxTwo.setReadOnly(True)
        self.h3.addWidget(self.myTextBoxTwo,4)
        
        # self.lineOne = QLabel('/'*250, self) 
        # self.lineOne.setMaximumWidth(800)
        # self.v1.addWidget(self.lineOne)
        
        self.CheckOne = QCheckBox('Abrir de inmediato el documento generado', self)  
        self.CheckOne.setFont(fontTwo)
        self.CheckOne.setMinimumHeight(35)
        # self.CheckOne.setMaximumWidth(800)
        self.CheckOne.setStyleSheet("QCheckBox {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255);padding-left:10px;}") 
        self.CheckOne.setChecked(True)
        self.v1.addWidget(self.CheckOne)
             
        # self.lineTwo = QLabel('/'*250, self)
        # self.lineTwo.setMaximumWidth(800)
        # self.v1.addWidget(self.lineTwo)     
        self.h2.addStretch()
        self.start = QPushButton('Ejecutar', self)
        self.start.setStyleSheet(self.style1)
        self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        # self.start.setMaximumWidth(200)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
    
        self.button = QPushButton('Limpiar', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMaximumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
        
        
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        
        self.effect = QGraphicsOpacityEffect(self)
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 50, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)       
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)
        # self.mainLayout.setSpacing(30)
        # self.v1.setSpacing(0)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.h3)
        self.mainLayout.addLayout(self.v1)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)       
        self.setLayout(self.mainLayout)
        
        quit = QAction("Quit", self)
        quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            if self.var1 is not None and self.var2 is not None:
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.state = self.CheckOne.checkState()
                self.threadpool = QThreadPool()
                self.runner = JobRunnerFour(self.var1,self.var2,self.state)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Intenta de nuevo.')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.myTextBoxTwo.setText(None)
        self.var1=None
        self.var2=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.progress.hide()
        
    def openFileNameDialogOne(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileNames, _ = QFileDialog.getOpenFileNames(self,"Selecciona tus imágenes","","Imágenes (*.png *.jpg *.jpeg)",options=options)
        files=[]
        if fileNames: 
            for fileName in fileNames:
                fileName=os.path.abspath(fileName)
                files.append(fileName)
            self.myTextBoxOne.setText(str(files).strip('[').strip(']'))
            self.var1=files
            
        return files
    def openFileNameDialogTwo(self):
    
        fileName, _ = QFileDialog.getSaveFileName(self,"Guardar como",'',filter="PDF (*.pdf)")
        
        if fileName:        
            if '.pdf' not in fileName:
                fileName=fileName+'.pdf'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxTwo.setText(fileName)
            self.var2=self.myTextBoxTwo.text()
        return fileName
  
    def alert(self, msg):
        if msg=='Error1':
            self.error('No puedes guardar el resultado encima de uno de los PDFs que deseas unir.')
        elif msg=='Error2':
            self.error('Cierra el PDF sobre el cual deseas guardar el resultado.')
        self.clean()
        
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.myTextBoxTwo.setText(None)
            self.var1=None
            self.var2=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('¡Listo, ya puedes visualizar tus documentos!')
            self.labelThree.show()
            
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()

    def closeEvent(self, event):
        close = QMessageBox()
        # close.setWindowTitle(self.title)
        close.setWindowTitle("¿Estás seguro?")
        close.setWindowIcon(QIcon(icon))
        close.setFont(fontTwo)
        close.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        # close.setText("¿Estás seguro?")
        # close.setInformativeText('Se detendrá la función si se está ejecutando, pero no te preocupes ya que se guardará el avance.')
        close.setText("¿Estás seguro que deseas salir?")
        close.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        close = close.exec()

        if close == QMessageBox.Yes:           
            event.accept()  
            self.clean()
        else:
            event.ignore()
    
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None
    
    def instructions(self):
        info = QMessageBox()
        info.setWindowTitle(choices[3][3:])
        
        info.setWindowIcon(QIcon(icon))
        info.setText('work in progress')

        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setText('Entendido')
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()
  
# <codecell>    
class WorkerSignalsFive(QObject):
  
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    
class JobRunnerFive(QRunnable):    
    signals = WorkerSignalsFive()
    
    def __init__(self,SaveAs,name,choice):
        super().__init__()

        self.is_killed = False 
        self.SaveAs=SaveAs
        self.name=name
        self.choice=choice
        
    @pyqtSlot()
        
    def run(self):
        
        
        RDDir=os.path.abspath(os.path.dirname(self.SaveAs))
        if self.name in os.listdir(RDDir):
            self.signals.alert.emit('Error1')
        else:   
            if self.choice==0:
                doc = fitz.open(self.SaveAs)
                page = doc.loadPage(0)              
                text = page.getText()  
                nro_req=findall('(?<!\d)\d{13}(?!\d)', text) 
                nro_ruc=findall('(?<!\d)\d{11}(?!\d)', text)  
                doc.close()
                if len(nro_ruc)==0 or len(nro_req)==0:
                    self.signals.alert.emit('Error2')
                else:
                    os.mkdir(RDDir+'\\'+self.name)          
                    targetOne=RDDir+'\\'+self.name
                    os.mkdir(targetOne+'\\'+self.name)
                    targetTwo=RDDir+'\\'+self.name+'\\'+self.name
                    
                    nro_ruc = nro_ruc[0]
                    nro_req = nro_req[0]
           
                    text_RD_path=targetTwo+'\\files.txt'
                    with open(text_RD_path, 'w') as text_RD:
                        text_RD.write(nro_ruc+'|'+nro_req)
                    
                    shutil.copy(self.SaveAs,targetTwo+'\\'+nro_ruc+'_'+nro_req+'.pdf')
                    shutil.make_archive(targetTwo, 'zip',root_dir=targetOne , base_dir=self.name)
                    
                    time.sleep(1)
                    self.signals.finished.emit('Done')
            elif self.choice==1:
                carta=0
                doc=fitz.open(self.SaveAs)
                page = doc.loadPage(0)              
                text = page.getText() 
                doc.close()
                
                lines=text.split('\n')
                for line in lines:
                    if ('CARTA' or 'Carta' or 'carta') in line:
                        target=line
                        target=target.replace(' ','')
                        if '/' in target:
                            carta=1                  
                        target=target.split('-')
                        break
                if carta==1: 
                    OF=findall('(?<!\d)\d{12}(?!\d)', text)
                    num=findall('[0-9]+', target[0])
                    if len(num)==0 or len(OF)==0:
                        self.signals.alert.emit('Error3')
                    else:
                        os.mkdir(RDDir+'\\'+self.name)          
                        targetOne=RDDir+'\\'+self.name
                        os.mkdir(targetOne+'\\'+self.name)
                        targetTwo=RDDir+'\\'+self.name+'\\'+self.name
                        
                        num=num[0]
                        OF=OF[0]
                        there=len(num)
                        added=(7-there)*'0'
                        num=added+num+target[1]
                        area=target[2].split('/')[1]
                        num=num+area
                        
                        ruc=findall('(?<!\d)\d{11}(?!\d)', text)
                        ruc=ruc[0]
                        
                        text_RD_path=targetTwo+'\\files.txt'
                        with open(text_RD_path, 'w') as text_RD:
                            text_RD.write(ruc+'|'+num+'|'+OF)
                        
                        shutil.copy(self.SaveAs,targetTwo+'\\'+ruc+'_'+num+'.pdf')
                        shutil.make_archive(targetTwo, 'zip',root_dir=targetOne , base_dir=self.name)
                        
                        time.sleep(1)
                        self.signals.finished.emit('Done')
                else:               
                    OF=findall('[0-9]+', target[0])
                    ruc=findall('(?<!\d)\d{11}(?!\d)', text)
                    if len(OF)==0 or len(ruc)==0:
                        self.signals.alert.emit('Error4')
                    else:
                        os.mkdir(RDDir+'\\'+self.name)          
                        targetOne=RDDir+'\\'+self.name
                        os.mkdir(targetOne+'\\'+self.name)
                        targetTwo=RDDir+'\\'+self.name+'\\'+self.name
                        
                        OF=OF[0]
                        num=OF+target[1]
                        ruc=ruc[0]
                        
                        text_RD_path=targetTwo+'\\files.txt'
                        with open(text_RD_path, 'w') as text_RD:
                            text_RD.write(ruc+'|'+num+'|'+OF)
                        shutil.copy(self.SaveAs,targetTwo+'\\'+ruc+'_'+num+'.pdf')
                        shutil.make_archive(targetTwo, 'zip',root_dir=targetOne , base_dir=self.name)
                        
                        time.sleep(1)
                        self.signals.finished.emit('Done')
                    
    def kill(self):
        self.is_killed = True
           
class ActionsFive(QWidget):
    """
    Simple dialog that consists of a Progress Bar and a Button.
    Clicking on the button results in the start of a timer and
    updates the progress bar.
    """
 
    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'       
        self.var1=None
        self.initUI()
        self.msg1='Verifica los datos ingresados.'
        
    def initUI(self):
        self.style = QApplication.style()
       
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}") 
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
        self.style4=("QComboBox {selection-background-color: rgb(69, 70, 77);background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);padding-left:10px}"
                     "QComboBox QAbstractItemView::item { min-height: 35px; min-width: 50px;}"
                     "QListView::item { color: white; background-color: rgb(69, 70, 77)}"
                     "QListView::item:selected { color: white; background-color: IndianRed}") 
        self.style5=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}")                     
         
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.h3=QHBoxLayout()
        self.h4=QHBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))
    
      
        self.buttonTwo = QPushButton('Cargar PDF', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h1.addWidget(self.myTextBoxOne,4)
        
        self.buttonThree = QPushButton('Nombre', self)  
        self.buttonThree.setMinimumHeight(35)
        # self.buttonThree.setMaximumWidth(200)
        self.buttonThree.setStyleSheet(self.style5)
        self.buttonThree.setFont(fontTwo)
        self.buttonThree.setEnabled(False)
        self.h3.addWidget(self.buttonThree,1)
        
        self.myTextBoxTwo = QLineEdit(self)
        self.myTextBoxTwo.setMinimumHeight(35)  
        self.myTextBoxTwo.setStyleSheet('background-color: rgb(69, 70, 77); color: white')  
        self.myTextBoxTwo.setFont(fontTwo)
        self.myTextBoxTwo.setPlaceholderText('Ingresa un nombre para el archivo .zip')
        self.myTextBoxTwo.setToolTip('Escribe un nombre para la carpeta que contendrá el documento')
        self.h3.addWidget(self.myTextBoxTwo,4)
                                  
        self.buttonFour = QPushButton('Opciones', self)  
        self.buttonFour.setMinimumHeight(35)  
        # self.buttonFour.setMaximumWidth(200)
        self.buttonFour.setFont(fontTwo)
        self.buttonFour.setStyleSheet(self.style5)
        self.buttonFour.setEnabled(False)
        self.h4.addWidget(self.buttonFour,1)
        
        self.combo=QComboBox(self)
        self.combo.addItems(['Requerimiento y resultado',
                             'Cartas'])
        self.combo.setMinimumHeight(35)  
        # self.combo.setMaximumWidth(600)
        self.combo.setFont(fontTwo)
        self.combo.setStyleSheet(self.style4)
        self.listview=QListView()
        self.listview.setFont(fontTwo)
        self.listview.setCursor(QCursor(Qt.PointingHandCursor))
        self.combo.setView(self.listview)
        self.combo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h4.addWidget(self.combo,4)
        
        self.h2.addStretch()             
        self.start = QPushButton('Ejecutar', self)
        self.start.setStyleSheet(self.style1)
        self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        # self.start.setMaximumWidth(200)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
    
        self.button = QPushButton('Limpiar', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMaximumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()       
        
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        
        self.effect = QGraphicsOpacityEffect(self)
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 50, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)       
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)
        # self.mainLayout.setSpacing(30)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.h3)
        self.mainLayout.addLayout(self.h4)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)       
        self.setLayout(self.mainLayout)
        
        quit = QAction("Quit", self)
        quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            self.name=self.myTextBoxTwo.text().strip()
            if self.var1 is not None and len(self.name.strip())>0:
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.name=self.myTextBoxTwo.text()
                self.choice = self.combo.currentIndex()
                print(self.choice)
                self.threadpool = QThreadPool()
                self.runner = JobRunnerFive(self.var1,self.name,self.choice)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Intenta de nuevo.')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.var1=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.myTextBoxTwo.setText('')
        self.progress.hide()
        self.runner=None

    def openFileNameDialogOne(self):

        fileName, _ = QFileDialog.getOpenFileName(self,"Carga tu documento",'',filter="PDF (*.pdf)")
        
        if fileName:        
            if '.pdf' not in fileName:
                fileName=fileName+'.pdf'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxOne.setText(fileName)
            self.var1=self.myTextBoxOne.text()
        return fileName
  
    def alert(self, msg):
        if msg=='Error1':
            self.error('Ya existe una carpeta con ese nombre.')
        elif msg=='Error2':
            self.error('No se encontró el número de RUC y/o Requerimiento.')
        elif msg=='Error3':
            self.error('No se encontró el número de Carta y/o OF.')
        elif msg=='Error4':
            self.error('No se encontró el número de RUC y/o OF.')
        self.clean()
        
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.var1=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('¡Listo, ya puedes visualizar tus documentos!')
            self.labelThree.show()
            
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()

    def closeEvent(self, event):
        close = QMessageBox()
        # close.setWindowTitle(self.title)
        close.setWindowTitle("¿Estás seguro?")
        close.setWindowIcon(QIcon(icon))
        close.setFont(fontTwo)
        close.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        # close.setText("¿Estás seguro?")
        # close.setInformativeText('Se detendrá la función si se está ejecutando, pero no te preocupes ya que se guardará el avance.')
        close.setText("¿Estás seguro que deseas salir?")
        close.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        close = close.exec()

        if close == QMessageBox.Yes:           
            event.accept()        
            self.clean()
        else:
            event.ignore()
   
    
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
    
    def instructions(self):
        info = QMessageBox()
        info.setWindowTitle(choices[4][3:])
        
        info.setWindowIcon(QIcon(icon))
        info.setText('''El documento resultante se guardará en la carpeta que contiene tu Requerimiento/Carta''')
        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setText('Entendido')
        buttonOk.setFont(fontOne)
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()
# <codecell>  
class WorkerSignalsSix(QObject):
    alert=pyqtSignal(str)
    finished=pyqtSignal(str)
    reportMsg=pyqtSignal(str)
    
class JobRunnerSix(QRunnable):    
    signals = WorkerSignalsSix()
    
    def __init__(self,rd,rm,rpv):
        super().__init__()

        self.is_killed = False 
        self.rd=rd
        self.rm=rm
        self.rpv=rpv
        
    @pyqtSlot()

    def run(self):
        

        text=''
        doc = fitz.open(self.rpv)
        for i in range(len(doc)):
            page = doc.loadPage(i) 
            text += page.getText() 
        doc.close()
        
        ruc=findall('\\b2\d{10}\\b', text)
        ruc=ruc[0]
        pre=findall('\d{15}', text)
        tipo=findall('RD|RM', text)
        targetName='Valores '+ruc
        names=[]
        for p in pre:
            names.append(ruc+'_'+p+'_'+'01')
        
        if 'RM' in tipo and self.rm is None:
            self.signals.alert.emit('rm')
        elif 'RD' in tipo and self.rd is None:
            self.signals.alert.emit('rd') 
        else:
        
            RDDir=os.path.abspath(os.path.dirname(self.rd))
            if targetName in os.listdir(RDDir):
                self.signals.alert.emit('Error')
            else:
                os.mkdir(RDDir+'\\'+targetName)          
                targetOne=RDDir+'\\'+targetName
                os.mkdir(targetOne+'\\'+targetName)
                targetTwo=RDDir+'\\'+targetName+'\\'+targetName
                
                for i in range(len(names)):
                    if tipo[i]=='RD':
                        shutil.copy(self.rd,targetTwo+'\\'+names[i]+'.pdf')
                    elif tipo[i]=='RM':
                        shutil.copy(self.rm,targetTwo+'\\'+names[i]+'.pdf')
                
                self.signals.reportMsg.emit('Se encontraron '+str(tipo.count('RD'))+' RDs y '+str(tipo.count('RD'))+' RMs')
                
                text_RD_path=targetTwo+'\\files.txt'
                with open(text_RD_path, 'w') as text_RD:
                    for i in range(len(names)):
                        text_RD.write(ruc+'|'+pre[i]+'|'+'01\n')
             
                shutil.make_archive(targetTwo, 'zip',root_dir=targetOne , base_dir=targetName)
                time.sleep(3)
                self.signals.reportMsg.emit('Se encontraron '+str(tipo.count('RD'))+' RDs y '+str(tipo.count('RD'))+' RMs')

                self.signals.finished.emit('Done')
              
                
    def kill(self):
        self.is_killed = True
           
class ActionsSix(QWidget):
    """
    Simple dialog that consists of a Progress Bar and a Button.
    Clicking on the button results in the start of a timer and
    updates the progress bar.
    """
 
    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.var2=None
        self.var3=None
        self.initUI()
        self.msg1='Verifica los datos ingresados.' 
        
    def initUI(self):
        self.style = QApplication.style()
       
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}") 
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
        
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.h3=QHBoxLayout()
        self.h4=QHBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))
    
    
        self.buttonTwo = QPushButton('RD', self)   
        self.buttonTwo.clicked.connect(self.openFileNameDialogOne)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h1.addWidget(self.buttonTwo,1)
        
        self.buttonThree = QPushButton('RM', self)      
        self.buttonThree.clicked.connect(self.openFileNameDialogTwo)
        self.buttonThree.setMinimumHeight(35)
        # self.buttonThree.setMaximumWidth(200)
        self.buttonThree.setStyleSheet(self.style2)
        self.buttonThree.setFont(fontTwo)
        self.buttonThree.setCursor(QCursor(Qt.PointingHandCursor)) 
        self.h3.addWidget(self.buttonThree,1)
        
        self.buttonFour = QPushButton('RPV', self)      
        self.buttonFour.clicked.connect(self.openFileNameDialogThree)
        self.buttonFour.setMinimumHeight(35)
        # self.buttonFour.setMaximumWidth(200)
        self.buttonFour.setStyleSheet(self.style2)
        self.buttonFour.setFont(fontTwo)
        self.buttonFour.setCursor(QCursor(Qt.PointingHandCursor))
        self.h4.addWidget(self.buttonFour,1)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.myTextBoxOne.setPlaceholderText('Déjalo en blanco si no tienes RDs')
        self.h1.addWidget(self.myTextBoxOne,4)
        
        self.myTextBoxTwo = QLineEdit(self)
        self.myTextBoxTwo.setMinimumHeight(35)  
        self.myTextBoxTwo.setStyleSheet('background-color: rgb(69, 70, 77); color: white')  
        self.myTextBoxTwo.setFont(fontTwo)
        self.myTextBoxTwo.setReadOnly(True)
        self.myTextBoxTwo.setPlaceholderText('Déjalo en blanco si no tienes RMs')
        self.h3.addWidget(self.myTextBoxTwo,4)
        
        self.myTextBoxThree = QLineEdit(self)
        self.myTextBoxThree.setMinimumHeight(35)  
        self.myTextBoxThree.setStyleSheet('background-color: rgb(69, 70, 77); color: white')    
        self.myTextBoxThree.setFont(fontTwo)
        self.myTextBoxThree.setReadOnly(True)
        self.h4.addWidget(self.myTextBoxThree,4)
         
        self.h2.addStretch()                   
        self.start = QPushButton('Ejecutar', self)
        self.start.setStyleSheet(self.style1)
        self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        # self.start.setMaximumWidth(200)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
    
        self.button = QPushButton('Limpiar', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMaximumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
        
        self.labelOne = QLabel('', self)
        self.labelOne.setFont(fontThree)
        self.labelOne.setAlignment(Qt.AlignCenter)
        self.labelOne.hide()
        
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()      

        self.effect = QGraphicsOpacityEffect(self)
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 50, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)
       
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)
        # self.mainLayout.setSpacing(30)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.h3)
        self.mainLayout.addLayout(self.h4)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelOne)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)       
        self.setLayout(self.mainLayout)
        
        quit = QAction("Quit", self)
        quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            if self.var3 is not None:
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.threadpool = QThreadPool()
                self.runner = JobRunnerSix(self.var1,self.var2,self.var3)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.alert.disconnect(self.alert)
                    self.runner.signals.finished.disconnect(self.finished)
                    self.runner.signals.reportMsg.disconnect(self.report)
                except TypeError:     
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                    self.runner.signals.reportMsg.connect(self.report)
                else:
                    self.runner.signals.alert.connect(self.alert)
                    self.runner.signals.finished.connect(self.finished)
                    self.runner.signals.reportMsg.connect(self.report)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Intenta de nuevo.')
                self.error(self.msg1)
                self.progress.hide()
                
    def clean(self):
        
        self.myTextBoxOne.setText(None)
        self.myTextBoxTwo.setText(None)
        self.myTextBoxThree.setText(None)
        self.var1=None
        self.var2=None
        self.var3=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelOne.setText('')
        self.labelThree.hide()
        self.progress.hide()

    def openFileNameDialogOne(self):

        fileName, _ = QFileDialog.getOpenFileName(self,"Carga tu RD",'',"PDF (*.pdf)")
       
        if fileName:   
            fileName=os.path.abspath(fileName)
            # print(fileName)
            self.myTextBoxOne.setText(fileName)
            self.var1=self.myTextBoxOne.text()
            
        return fileName
    def openFileNameDialogTwo(self):

        fileName, _ = QFileDialog.getOpenFileName(self,"Carga tu RM",'',"PDF (*.pdf)")
        
        if fileName:        
            fileName=os.path.abspath(fileName)
            # print(fileName)
            self.myTextBoxTwo.setText(fileName)
            self.var2=self.myTextBoxTwo.text()
        return fileName
    def openFileNameDialogThree(self):

        fileName, _ = QFileDialog.getOpenFileName(self,"Carga tu RPV",'',"PDF (*.pdf)")
        
        if fileName:        
            fileName=os.path.abspath(fileName)
            # print(fileName)
            self.myTextBoxThree.setText(fileName)
            self.var3=self.myTextBoxThree.text()
        return fileName 
  
    def alert(self, msg):
        if msg=='Error':
            self.error('Ya existe un folder con ese nombre')
        elif msg=='rm':
            self.error('No has cargado las RMs')
        elif msg=='rd':
            self.error('No has cargado las RDs')
        self.clean()
    def report(self,msg):
        self.labelOne.setText(msg)
        self.labelOne.show()
    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxOne.setText(None)
            self.myTextBoxTwo.setText(None)
            self.myTextBoxThree.setText(None)
            self.var1=None
            self.var2=None
            self.var3=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('¡Listo, ya puedes visualizar tus documentos!')
            self.labelThree.show()
            self.labelThree.setPixmap(self.pixmap) 
            self.progress.hide()

    def closeEvent(self, event):
        close = QMessageBox()
        # close.setWindowTitle(self.title)
        close.setWindowTitle("Seguro?")
        close.setWindowIcon(QIcon(icon))
        close.setFont(fontTwo)
        close.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        # close.setText("¿Estás seguro?")
        # close.setInformativeText('Se detendrá la función si se está ejecutando, pero no te preocupes ya que se guardará el avance.')
        close.setText("¿Estás seguro que deseas salir?")
        close.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        close = close.exec()

        if close == QMessageBox.Yes:           
            event.accept()      
            self.clean()
        else:
            event.ignore()  
    
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None
    
    def instructions(self):
        info = QMessageBox()
        info.setWindowTitle(choices[5][3:])
        
        info.setWindowIcon(QIcon(icon))
        info.setText('''El documento resultante se guardará en la carpeta que contiene tu RD, bajo un nombre de la forma "Valores + RUC".''')
        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setText('Entendido')
        buttonOk.setFont(fontOne)
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()
 
# <codecell>  
class WorkerSignalsSeven(QObject):

    finished=pyqtSignal(str)
    
class JobRunnerSeven(QRunnable):    
    signals = WorkerSignalsSeven()
    
    def __init__(self,SaveAs,RUC,doc,profile):
        super().__init__()

        self.is_killed = False 
        self.SaveAs=SaveAs
        self.RUC=RUC
        if profile==0:
            self.profile='F01'
        elif profile==1:
            self.profile='F02'
        elif profile==2:
            self.profile='F03'
        elif profile==3:
            self.profile='F04'           
        self.doc=doc
        
    @pyqtSlot()
        
    def run(self):

        text='6,'+self.RUC+','+self.profile+','+self.doc+',DESCARGA LE'
        
        with open(self.SaveAs, "w") as text_file:
            text_file.write('')
            text_file.write(text)
        time.sleep(1)
        self.signals.finished.emit('Done')
                             
    def kill(self):
        self.is_killed = True   
        
class ActionsSeven(QWidget):
    """
    Simple dialog that consists of a Progress Bar and a Button.
    Clicking on the button results in the start of a timer and
    updates the progress bar.
    """
 
    def __init__(self):
        super().__init__()
        self.runner=None
        self.title = 'LuftMensch'
        self.var1=None
        self.msg1='El RUC debe tener 11 dígitos y el doc. de sustento 12 o 18.'
        self.initUI()
    def initUI(self):
        self.style = QApplication.style()
        
        
        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}")        
        self.style3 = ("QProgressBar {border: 2px solid grey;border-radius: 5px;text-align: center}"
                         "QProgressBar::chunk {background-color: IndianRed;width: 10px;margin: 1px;}")
        self.style4=("QComboBox {selection-background-color: rgb(69, 70, 77);background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);padding-left:10px}"
                     "QComboBox QAbstractItemView::item { min-height: 35px; min-width: 50px;}"
                     "QListView::item { color: white; background-color: rgb(69, 70, 77)}"
                     "QListView::item:selected { color: white; background-color: IndianRed}") 
        self.style5=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}")                     
         
        self.setWindowTitle(self.title)
             
        self.h1=QHBoxLayout()
        self.h2=QHBoxLayout()
        self.h3=QHBoxLayout()
        self.h4=QHBoxLayout()
        self.h5=QHBoxLayout()
        self.v2=QVBoxLayout()
                               
        # self.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(86, 88, 110)")
        self.setWindowIcon(QIcon(icon))
    
        self.buttonOne = QPushButton('RUC', self) 
        self.buttonOne.setMinimumHeight(35)
        # self.buttonOne.setMaximumWidth(200)
        self.buttonOne.setStyleSheet(self.style5)
        self.buttonOne.setFont(fontTwo)
        self.buttonOne.setEnabled(False)
        self.h1.addWidget(self.buttonOne,1)
        
        self.myTextBoxZero = QLineEdit(self)
        self.myTextBoxZero.setMinimumHeight(35)  
        self.myTextBoxZero.setStyleSheet('background-color: rgb(69, 70, 77); color: white')
        self.myTextBoxZero.setFont(fontTwo)
        self.myTextBoxZero.setPlaceholderText('Ingresa el número de RUC')
        self.h1.addWidget(self.myTextBoxZero,4)
        
        self.buttonThree = QPushButton('Doc. Sustento', self)  
        self.buttonThree.setMinimumHeight(35)
        # self.buttonThree.setMaximumWidth(200)
        self.buttonThree.setStyleSheet(self.style5)
        self.buttonThree.setFont(fontTwo)
        self.buttonThree.setEnabled(False)
        self.h3.addWidget(self.buttonThree,1)
        
        self.buttonTwo = QPushButton('Guardar como', self)      
        self.buttonTwo.clicked.connect(self.openFileNameDialogTwo)
        self.buttonTwo.setMinimumHeight(35)
        # self.buttonTwo.setMaximumWidth(200)
        self.buttonTwo.setStyleSheet(self.style2)
        self.buttonTwo.setFont(fontTwo)
        self.buttonTwo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h5.addWidget(self.buttonTwo,1)
        
        self.myTextBoxTwo = QLineEdit(self)
        self.myTextBoxTwo.setMinimumHeight(35)  
        self.myTextBoxTwo.setStyleSheet('background-color: rgb(69, 70, 77); color: white')  
        self.myTextBoxTwo.setFont(fontTwo)
        self.myTextBoxTwo.setPlaceholderText('Ingresa el número de OF o equivalente')
        self.h3.addWidget(self.myTextBoxTwo,4)
        
        self.myTextBoxOne = QLineEdit(self)
        self.myTextBoxOne.setMinimumHeight(35)  
        self.myTextBoxOne.setStyleSheet('background-color: rgb(69, 70, 77); color: white') 
        self.myTextBoxOne.setFont(fontTwo)
        self.myTextBoxOne.setReadOnly(True)
        self.h5.addWidget(self.myTextBoxOne,4)
                                  
        self.buttonFour = QPushButton('Perfiles', self)  
        self.buttonFour.setMinimumHeight(35)  
        # self.buttonFour.setMaximumWidth(200)
        self.buttonFour.setFont(fontTwo)
        self.buttonFour.setStyleSheet(self.style5)
        self.buttonFour.setEnabled(False)
        self.h4.addWidget(self.buttonFour,1)
        
        self.combo=QComboBox(self)
        self.combo.addItems(['F01 - Orden de Fiscalización',
       'F02 - Acción Inductiva - Esquela',
       'F03 - Programa de Fiscalización - Aduanas',
       'F04 - Acción Inductiva - Carta Inductiva'])
        self.combo.setMinimumHeight(35)  
        # self.combo.setMaximumWidth(600)
        self.combo.setFont(fontTwo)
        self.combo.setStyleSheet(self.style4)
        self.listview=QListView()
        self.listview.setFont(fontTwo)
        self.listview.setCursor(QCursor(Qt.PointingHandCursor))
        self.combo.setView(self.listview)
        self.combo.setCursor(QCursor(Qt.PointingHandCursor))
        self.h4.addWidget(self.combo,4)
        
        self.h2.addStretch()              
        self.start = QPushButton('Ejecutar', self)
        self.start.setStyleSheet(self.style1)
        self.start.setFocus()
        self.start.setFont(fontOne)
        self.start.setMinimumHeight(35)
        # self.start.setMaximumWidth(200)
        self.start.setEnabled(True)
        self.start.setCursor(QCursor(Qt.PointingHandCursor))
        self.start.clicked.connect(self.started) 
        self.h2.addWidget(self.start)
    
        self.button = QPushButton('Limpiar', self)
        self.button.setStyleSheet(self.style1)
        self.button.setFont(fontOne)
        self.button.setMinimumHeight(35)
        # self.button.setMaximumWidth(200)
        self.button.setEnabled(True)
        self.button.setCursor(QCursor(Qt.PointingHandCursor))
        self.button.clicked.connect(self.clean) 
        self.h2.addWidget(self.button)
        
        self.progress = QProgressBar(self)
        self.progress.setFormat("")
        self.progress.setStyleSheet(self.style3)    
        self.progress.setFont(fontOne)
        # self.progress.setMaximumWidth(800)
        self.progress.setAlignment(Qt.AlignCenter) 
        self.progress.setValue(0)
        self.progress.setMaximum(0)
        self.progress.hide()
        
        
        self.labelTwo = QLabel('', self)
        self.labelTwo.setFont(fontThree)
        self.labelTwo.setStyleSheet("color:LightGreen")
        self.labelTwo.setAlignment(Qt.AlignCenter)
        # self.labelTwo.hide()
        
        self.effect = QGraphicsOpacityEffect(self)
        self.pixmap = QPixmap(pic)
        self.pixmap = self.pixmap.scaled(50, 50, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setAlignment(Qt.AlignCenter)       
        # self.info.setIcon(QIcon(self.style.standardIcon(QStyle.SP_FileDialogInfoView)))  
        
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)
        # self.mainLayout.setSpacing(30)
        self.mainLayout.addLayout(self.h1)
        self.mainLayout.addLayout(self.h3)
        self.mainLayout.addLayout(self.h4)
        self.mainLayout.addLayout(self.h5)
        self.mainLayout.addLayout(self.h2)
        self.mainLayout.addWidget(self.progress)
        self.mainLayout.addWidget(self.labelTwo)
        self.mainLayout.addWidget(self.labelThree)       
        self.setLayout(self.mainLayout)
        
        quit = QAction("Quit", self)
        quit.triggered.connect(self.closeEvent)
   
    def started(self):
        
        if self.runner is None:
            self.start.setEnabled(False)
            self.RUC=self.myTextBoxZero.text().strip()
            self.doc=self.myTextBoxTwo.text().strip()
            if self.var1 is not None and len(self.RUC)==11 and (len(self.doc)==12 or len(self.doc)==18):
                self.labelTwo.setText('')
                self.labelThree.hide()
                self.progress.show()
                self.name=self.myTextBoxTwo.text()
                self.choice = self.combo.currentIndex()
                print(self.choice)
                self.threadpool = QThreadPool()
                self.runner = JobRunnerSeven(self.var1,self.RUC,self.doc,self.choice)   
                self.threadpool.start(self.runner)                                         
                try:
                    self.runner.signals.finished.disconnect(self.finished)
                except TypeError:     
                    self.runner.signals.finished.connect(self.finished)
                else:
                    self.runner.signals.finished.connect(self.finished)
            else:
                self.start.setEnabled(True)
                self.labelTwo.setText('Intenta de nuevo.')
                self.error(self.msg1)
                
    def clean(self):
        
        self.myTextBoxZero.setText(None)
        self.myTextBoxOne.setText(None)
        self.var1=None
        self.runner=None
        self.labelTwo.setText('')
        self.labelThree.hide()
        self.myTextBoxTwo.setText('')
        self.progress.hide()

    def openFileNameDialogTwo(self):
        
        fileName, _ = QFileDialog.getSaveFileName(self,"Guardar como",'',filter="Texto (*.txt)")
        
        if fileName:        
            if '.txt' not in fileName:
                fileName=fileName+'.txt'
            fileName=os.path.abspath(fileName)         
            self.myTextBoxOne.setText(fileName)
            self.var1=self.myTextBoxOne.text()
        return fileName

    def finished(self, msg):
        if msg=='Done':
            self.runner=None
            self.myTextBoxZero.setText(None)
            self.myTextBoxOne.setText(None)
            self.myTextBoxTwo.setText(None)
            self.var1=None
            self.start.setEnabled(True)   
            self.labelTwo.setText('¡Listo, ya puedes visualizar tus documentos!')
            self.labelThree.show()
            
            self.labelThree.setPixmap(self.pixmap) 
            self.labelThree.show()
            self.progress.hide()

    def closeEvent(self, event):
        close = QMessageBox()
        # close.setWindowTitle(self.title)
        close.setWindowTitle("¿Seguro?")
        close.setWindowIcon(QIcon(icon))
        close.setFont(fontTwo)
        close.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        # close.setText("¿Estás seguro?")
        # close.setInformativeText('Se detendrá la función si se está ejecutando, pero no te preocupes ya que se guardará el avance.')
        close.setText("¿Estás seguro que deseas salir?")
        close.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        close = close.exec()

        if close == QMessageBox.Yes:           
            event.accept()   
            self.clean()
        else:
            event.ignore()
     
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Verifica los datos ingresados")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()
        self.start.setEnabled(True)
        self.runner=None
    
    def instructions(self):
        info = QMessageBox()
        info.setWindowTitle(self.title)
        
        info.setWindowIcon(QIcon(icon))
        info.setText("Intrucciones de uso.")
        info.setInformativeText(
        '''
Estas son las instrucciones de uso de la opción 1.'''
        )
        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setText('Entendido')
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()
       
# <codecell>  
    
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        
        self.window1 = ActionsOne()
        self.window2 = ActionsTwo()
        self.window3 = ActionsThree()
        self.window4 = ActionsFour()
        self.window5 = ActionsFive()
        self.window6 = ActionsSix()
        self.window7 = ActionsSeven()   
        self.title = 'LuftMensch'
        self.initUI()
        
    def initUI(self):  

        self.style1=("QPushButton { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}"
                     "QPushButton:hover { background-color: rgba(155, 61, 61,230) ;color: white;}"
                      "QPushButton:pressed { background-color: rgb(69, 70, 77) ;color: rgb(255, 255, 255 );}")
        self.style2=("QPushButton { background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);}"
                      "QPushButton:hover { background-color: rgba(69, 70, 77,230) ;color: white;}"
                      "QPushButton:pressed { background-color: rgb(155, 61, 61 ); color: rgb(255, 255, 255 );}")
        self.style4=("QComboBox {selection-background-color: rgb(69, 70, 77);background-color: rgb(69, 70, 77); color: rgb(255, 255, 255);padding-left:10px}"
                     "QComboBox QAbstractItemView::item { min-height: 35px; min-width: 50px;}"
                     "QListView::item { color: white; background-color: rgb(69, 70, 77)}"
                     "QListView::item:selected { color: white; background-color: IndianRed}") 
        
        self.style = QApplication.style()
       
        self.setWindowTitle(self.title)       
        # self.setMinimumSize(750,500)
        self.setMinimumSize(520,520)
        # self.resize(500,600)
        # self.move(500, 2)
        # self.setWindowState(Qt.WindowMaximized)
        self.setStyleSheet("background-color: rgb(22, 23, 24); color:CornflowerBlue")
        self.setWindowIcon(QIcon(icon))
        
        self.menuBar = self.menuBar()
        self.menuBar.setCursor(QCursor(Qt.PointingHandCursor))
        self.menuBar.setStyleSheet("QMenuBar {background-color: rgb(155, 61, 61); color: rgb(255, 255, 255)}"
                                   "QMenuBar:item:selected {background-color: white ;color: black}") 
        self.menuBar.addAction('&Acerca de', self.about)
        self.menuBar.addAction('&Ir al repositorio', self.repo)
        self.menuBar.addAction('&Actualizar', self.update)
        # self.menuBar.addAction('&Ayuda', self.need_help)
        
        self.help=QMenu("&Instrucciones")
        self.help.setStyleSheet("QMenu {background-color: white; color: black}"
                                   "QMenu:item:selected {background-color: white ;color: rgb(155, 61, 61)}") 
        self.menuBar.addMenu(self.help)
        self.help.setCursor(QCursor(Qt.PointingHandCursor))
        self.help.addAction(choices[0], self.window1.instructions)
        self.help.addAction(choices[4], self.window5.instructions) 
        self.help.addAction(choices[5], self.window6.instructions)       
        self.stackedLayout = QStackedLayout()
              
        self.mainLayout = QVBoxLayout()
        self.mainLayout.setAlignment(Qt.AlignCenter)    
    
        self.h=QHBoxLayout()
        self.v=QVBoxLayout()
        
        self.v0=QVBoxLayout()
        self.v1=QVBoxLayout()
        self.v2=QVBoxLayout()
        self.v3=QVBoxLayout()
        self.h3=QHBoxLayout()
        self.h4=QHBoxLayout()       
        self.h5=QHBoxLayout()
        
        windows=[self.window1,
                 self.window2,
                 self.window3,
                 self.window4,
                 self.window5,
                 self.window6,
                 self.window7]    
        
        for window in windows:
            self.stackedLayout.addWidget(window)
            
        self.pageCombo = QComboBox()   
        self.pageCombo.addItems(choices)
        self.pageCombo.setMinimumHeight(35)
        self.pageCombo.setStyleSheet(self.style4)
        self.listview=QListView()
        self.listview.setFont(fontTwo)
        self.listview.setCursor(QCursor(Qt.PointingHandCursor))
        self.pageCombo.setView(self.listview)
        self.pageCombo.setCursor(QCursor(Qt.PointingHandCursor))
        self.pageCombo.setFont(fontTwo)
        self.pageCombo.activated.connect(self.toggle_window)

        self.v0.addWidget(self.pageCombo)

        self.h.addLayout(self.v1)
        self.h.addLayout(self.v)
        self.h.addLayout(self.v2)   

        self.stackedLayout.setAlignment(Qt.AlignCenter)
        self.h.setAlignment(Qt.AlignCenter)
               
        self.mainLayout.addLayout(self.h,1)   
        self.mainLayout.addLayout(self.v0,0)       
        self.mainLayout.addLayout(self.stackedLayout,4)        
      
        self.pixmap = QPixmap(icon)
        self.pixmap = self.pixmap.scaled(70, 70, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelThree = QLabel('', self)
        self.labelThree.setPixmap(self.pixmap) 
        self.labelThree.setAlignment(Qt.AlignCenter) 
        self.v1.addWidget(self.labelThree)
        
        self.logo = QPixmap(logo)
        self.logo = self.logo.scaled(110, 110, Qt.KeepAspectRatio,Qt.SmoothTransformation)
        self.labelFour = QLabel('', self)
        self.labelFour.setPixmap(self.logo) 
        self.labelFour.setAlignment(Qt.AlignCenter) 
        self.v.addWidget(self.labelFour)
        
        self.titleOne = QLabel('Versión 1.3.2', self)
        self.titleOne.setFont(fontFive)
        self.titleOne.setStyleSheet("color:	IndianRed")
        self.titleOne.setAlignment(Qt.AlignRight | Qt.AlignBottom)  
        self.v2.addWidget(self.titleOne)
        
        self.labelOne = QLabel('Hola, '+username, self)
        self.labelOne.setFont(fontFive)
        self.labelOne.setAlignment(Qt.AlignRight)  
        self.v2.addWidget(self.labelOne)        
        
        self.status_label = QLabel()
        self.statusBar().addPermanentWidget(self.status_label)
        self.status_label.setText('Estás usando la versión 1.3.2, lanzada en marzo del 2021.')

        self.w = QWidget(self)
        self.w.setLayout(self.mainLayout)
        self.setCentralWidget(self.w)
        
        quit = QAction("Quit", self)
        quit.triggered.connect(self.closeEvent)
        
    def toggle_window(self):
        self.stackedLayout.setCurrentIndex(self.pageCombo.currentIndex())
    def error(self,errorMsg):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Critical)
        msg.setWindowTitle(self.title)
        msg.setWindowIcon(QIcon(icon))
        msg.setText("Error")
        msg.setFont(fontTwo)
        msg.setStandardButtons(QMessageBox.Ok)
        buttonOk = msg.button(QMessageBox.Ok)
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        msg.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        msg.setInformativeText(errorMsg)
        msg.exec_()

    def closeEvent(self, event):
        close = QMessageBox()
        close.setWindowTitle("¿Estás seguro?")
        close.setWindowIcon(QIcon(icon))
        close.setFont(fontTwo)
        close.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        close.setText("Se abandonará por completo la aplicación.")
        close.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        buttonYes = close.button(QMessageBox.Yes)
        buttonYes.setCursor(QCursor(Qt.PointingHandCursor))
        buttonYes.setFont(fontOne)
        buttonYes.setText('Sí')
        buttonCancel = close.button(QMessageBox.Cancel)
        buttonCancel.setText('No')
        buttonCancel.setCursor(QCursor(Qt.PointingHandCursor))
        buttonCancel.setFont(fontOne)
        close = close.exec()

        if close == QMessageBox.Yes:  
            
            event.accept() 
        else:
            event.ignore()
    def repo(self):
        from webbrowser import open as op
        op('https://github.com/lheredias/Luftmensch')
    def about(self):
        info = QMessageBox()
        info.setWindowTitle("Acerca de LuftMensch")
        
        info.setWindowIcon(QIcon(icon))
        info.setText('''Luftmensch es una aplicación de productividad y de código abierto pensada en automatizar ciertas tareas administrativas.

Si deseas revisar el historial de versiones, realizar consultas, dejar un comentario o hacer una sugerencia, visita el repositorio.''')

        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setText('Entendido')
        buttonOk.setCursor(QCursor(Qt.PointingHandCursor))
        buttonOk.setFont(fontOne)
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        retval = info.exec_()    
    def need_help(self):
        info = QMessageBox()
        info.setWindowTitle(self.title)
        
        info.setWindowIcon(QIcon(icon))
        info.setText('rerfsdgd')
 
        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        # info.setModal(True)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Ok)
        buttonOk = info.button(QMessageBox.Ok)
        buttonOk.setFont(fontOne)    
        buttonOk.setText('Releases')
        info.setDefaultButton(QMessageBox.Ok)
        info.show()
        info.exec_()       
    def update(self):
        info = QMessageBox()
        info.setWindowTitle("¿Cómo actualizar LuftMensch?")
        
        info.setWindowIcon(QIcon(icon))
        info.setText('''Para actualizar la aplicación, dale click en Releases y, una vez que termine de cargar la página, descarga la versión más reciente y guárdala encima de la actual.''')

        info.setFont(fontTwo)
        info.setStyleSheet("color: rgb(255, 255, 255); background-color: rgb(69, 70, 77  )")
        info.setWindowModality(0)
        info.activateWindow()
        info.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        buttonYes = info.button(QMessageBox.Yes)
        buttonYes.setCursor(QCursor(Qt.PointingHandCursor))
        buttonYes.setText('Releases')
        buttonYes.setFont(fontOne)
        buttonCancel = info.button(QMessageBox.Cancel)
        buttonCancel.setCursor(QCursor(Qt.PointingHandCursor))
        buttonCancel.setText('Entendido')
        buttonCancel.setFont(fontOne)
        info.setDefaultButton(QMessageBox.Cancel)
        info.show()
        retval = info.exec_()
        print(retval)
        if retval==16384:
            from webbrowser import open as op
            op('https://github.com/lheredias/Luftmensch/releases')
    
if __name__ == '__main__':
    os.environ["QT_AUTO_SCREEN_SCALE_FACTOR"] = "1"
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    app.setAttribute(Qt.AA_EnableHighDpiScaling,True)
    app.setAttribute(Qt.AA_UseHighDpiPixmaps, True)
    app.setWindowIcon(QIcon(icon))
    w = MainWindow()
    w.show() 
    sys.exit(app.exec_())
